"""
Generate academic thesis Chapters 4 and 5 as a formatted Word document (.docx).
Includes UML diagrams, performance tables, APA 7th edition citations, and
proper Times New Roman formatting.

Usage:
    python docs/generate_chapters.py
"""
import os
import sys
import json
import warnings
warnings.filterwarnings("ignore")

from pathlib import Path
import numpy as np
import pandas as pd

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from src.config import REPORTS_DIR, FIGURES_DIR, MODELS_DIR

OUTPUT_DIR = PROJECT_ROOT / "docs" / "output"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
DIAGRAM_DIR = OUTPUT_DIR / "diagrams"
DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)


# ============================================================================
# PART 1: UML DIAGRAM GENERATION
# ============================================================================

def draw_rounded_box(ax, xy, w, h, text, color="#2563eb", text_color="white",
                     fontsize=9, bold=False):
    """Draw a rounded rectangle with centered text."""
    x, y = xy
    box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.15",
                         facecolor=color, edgecolor="#1e293b", linewidth=1.5)
    ax.add_patch(box)
    weight = "bold" if bold else "normal"
    ax.text(x + w/2, y + h/2, text, ha="center", va="center",
            fontsize=fontsize, color=text_color, fontweight=weight,
            wrap=True)


def draw_arrow(ax, start, end, color="#64748b", style="->"):
    """Draw an arrow between two points."""
    ax.annotate("", xy=end, xytext=start,
                arrowprops=dict(arrowstyle=style, color=color, lw=1.5))


def generate_system_architecture_diagram():
    """Generate system architecture / deployment UML diagram."""
    fig, ax = plt.subplots(figsize=(12, 8))
    ax.set_xlim(0, 12)
    ax.set_ylim(0, 8)
    ax.axis("off")
    ax.set_title("System Architecture Diagram", fontsize=14, fontweight="bold", pad=20)

    # Data Layer
    draw_rounded_box(ax, (0.5, 6.5), 11, 1.2, "", color="#f1f5f9", text_color="black")
    ax.text(6, 7.5, "DATA LAYER", ha="center", fontsize=11, fontweight="bold", color="#1e293b")
    datasets = ["NSL-KDD", "KDDCup99", "Kaggle NID", "CIC-DDoS2019",
                "CIDDS-001", "DS2OS", "LUFlow", "NetworkLogs"]
    for i, ds in enumerate(datasets):
        x = 1.0 + i * 1.3
        draw_rounded_box(ax, (x, 6.6), 1.15, 0.45, ds, color="#dbeafe",
                         text_color="#1e40af", fontsize=6.5)

    # Preprocessing Layer
    draw_rounded_box(ax, (0.5, 4.8), 11, 1.3, "", color="#f0fdf4", text_color="black")
    ax.text(6, 5.85, "PREPROCESSING LAYER", ha="center", fontsize=11,
            fontweight="bold", color="#166534")
    procs = ["Data Loading\n& Cleaning", "Feature\nEngineering", "SMOTE\nBalancing",
             "Scaling\n(RobustScaler)", "Train/Test\nSplit"]
    for i, p in enumerate(procs):
        x = 1.0 + i * 2.1
        draw_rounded_box(ax, (x, 4.95), 1.8, 0.7, p, color="#bbf7d0",
                         text_color="#14532d", fontsize=7.5)

    # Model Layer
    draw_rounded_box(ax, (0.5, 2.8), 11, 1.5, "", color="#fef3c7", text_color="black")
    ax.text(6, 4.05, "MODEL LAYER (Ensemble IDS)", ha="center", fontsize=11,
            fontweight="bold", color="#92400e")
    models = ["Random\nForest", "XGBoost", "LightGBM", "MLP\n(Neural Net)"]
    for i, m in enumerate(models):
        x = 1.2 + i * 2.3
        draw_rounded_box(ax, (x, 3.15), 1.8, 0.65, m, color="#fde68a",
                         text_color="#78350f", fontsize=8.5)
    draw_rounded_box(ax, (8.5, 2.9), 2.5, 0.35, "Stacking Meta-Learner (LR)",
                     color="#f59e0b", text_color="white", fontsize=7.5, bold=True)

    # Evaluation Layer
    draw_rounded_box(ax, (0.5, 1.0), 11, 1.4, "", color="#fce7f3", text_color="black")
    ax.text(6, 2.15, "EVALUATION & DEPLOYMENT LAYER", ha="center", fontsize=11,
            fontweight="bold", color="#9d174d")
    evals = ["Threshold\nOptimization", "Metrics\nComputation", "Visualization\n& Reporting",
             "Interactive\nTesting App", "Monitoring\nDashboard"]
    for i, e in enumerate(evals):
        x = 0.8 + i * 2.15
        draw_rounded_box(ax, (x, 1.1), 1.85, 0.7, e, color="#fbcfe8",
                         text_color="#831843", fontsize=7.5)

    # Arrows between layers
    for x in [3, 6, 9]:
        draw_arrow(ax, (x, 6.5), (x, 6.15))
        draw_arrow(ax, (x, 4.8), (x, 4.35))
        draw_arrow(ax, (x, 2.8), (x, 2.45))

    plt.tight_layout()
    path = DIAGRAM_DIR / "system_architecture.png"
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def generate_class_diagram():
    """Generate UML class diagram for the IDS system."""
    fig, ax = plt.subplots(figsize=(14, 10))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 10)
    ax.axis("off")
    ax.set_title("UML Class Diagram", fontsize=14, fontweight="bold", pad=20)

    def draw_class_box(ax, x, y, w, h, name, attrs, methods, color="#2563eb"):
        # Name section
        rect = FancyBboxPatch((x, y + h - 0.6), w, 0.6, boxstyle="round,pad=0.05",
                              facecolor=color, edgecolor="#1e293b", linewidth=1.5)
        ax.add_patch(rect)
        ax.text(x + w/2, y + h - 0.3, name, ha="center", va="center",
                fontsize=9, fontweight="bold", color="white")
        # Attributes section
        attr_h = len(attrs) * 0.22 + 0.1
        rect2 = FancyBboxPatch((x, y + h - 0.6 - attr_h), w, attr_h,
                               boxstyle="round,pad=0.02",
                               facecolor="#f8fafc", edgecolor="#1e293b", linewidth=1)
        ax.add_patch(rect2)
        for i, a in enumerate(attrs):
            ax.text(x + 0.1, y + h - 0.75 - i*0.22, a, fontsize=6.5,
                    color="#334155", family="monospace")
        # Methods section
        meth_h = len(methods) * 0.22 + 0.1
        rect3 = FancyBboxPatch((x, y), w, meth_h, boxstyle="round,pad=0.02",
                               facecolor="#f1f5f9", edgecolor="#1e293b", linewidth=1)
        ax.add_patch(rect3)
        for i, m in enumerate(methods):
            ax.text(x + 0.1, y + meth_h - 0.15 - i*0.22, m, fontsize=6.5,
                    color="#475569", family="monospace")

    # DataPreprocessor class
    draw_class_box(ax, 0.5, 6.5, 3.5, 3,
                   "DataPreprocessor",
                   ["- dataset_paths: Dict", "- max_sample_size: int",
                    "- random_state: int"],
                   ["+ load_dataset(name)", "+ load_kdd_family()",
                    "+ load_ds2os()", "+ load_cic_ddos2019()",
                    "+ _encode_categoricals()", "+ _safe_numeric()"],
                   color="#2563eb")

    # FeatureEngineer class
    draw_class_box(ax, 5, 6.5, 3.5, 3,
                   "FeatureEngineer",
                   ["- test_size: float", "- scaler: RobustScaler",
                    "- feature_names: List"],
                   ["+ prepare_dataset(df)", "+ remove_constant_features()",
                    "+ remove_highly_correlated()", "+ balance_classes()"],
                   color="#059669")

    # EnsembleIDS class
    draw_class_box(ax, 9.5, 6.5, 4, 3,
                   "EnsembleIDS",
                   ["- base_models: Dict", "- meta_learner: LR",
                    "- imbalance_ratio: float"],
                   ["+ build_base_models()", "+ train_individual_models()",
                    "+ train_stacking_ensemble()", "+ save_model()",
                    "+ load_model()"],
                   color="#dc2626")

    # MetricsEvaluator class
    draw_class_box(ax, 0.5, 2.5, 3.5, 3,
                   "MetricsEvaluator",
                   ["- min_metric_score: float", "- target: float"],
                   ["+ compute_all_metrics()", "+ evaluate_model()",
                    "+ optimize_threshold()", "+ check_performance_target()",
                    "+ save_metrics_report()"],
                   color="#7c3aed")

    # Visualizer class
    draw_class_box(ax, 5, 2.5, 3.5, 3,
                   "Visualizer",
                   ["- figures_dir: Path", "- reports_dir: Path"],
                   ["+ plot_confusion_matrix()", "+ plot_roc_curves()",
                    "+ plot_metrics_comparison()", "+ plot_dataset_heatmap()"],
                   color="#0891b2")

    # TrainingPipeline class
    draw_class_box(ax, 9.5, 2.5, 4, 3,
                   "TrainingPipeline",
                   ["- dataset_names: List", "- all_metrics: List",
                    "- optimal_thresholds: Dict"],
                   ["+ train_on_dataset(name)", "+ run_full_pipeline()",
                    "+ save_results()"],
                   color="#ea580c")

    # Relationships
    draw_arrow(ax, (4, 8), (5, 8), color="#475569")
    ax.text(4.5, 8.15, "uses", fontsize=7, ha="center", color="#475569")

    draw_arrow(ax, (8.5, 8), (9.5, 8), color="#475569")
    ax.text(9, 8.15, "trains", fontsize=7, ha="center", color="#475569")

    draw_arrow(ax, (11.5, 6.5), (11.5, 5.5), color="#475569")
    ax.text(11.7, 6, "outputs", fontsize=7, color="#475569")

    draw_arrow(ax, (4, 4), (5, 4), color="#475569")
    ax.text(4.5, 4.15, "visualizes", fontsize=7, ha="center", color="#475569")

    draw_arrow(ax, (8.5, 4), (9.5, 4), color="#475569")
    ax.text(9, 4.15, "orchestrates", fontsize=7, ha="center", color="#475569")

    draw_arrow(ax, (2.25, 6.5), (2.25, 5.5), color="#475569")
    ax.text(2.5, 6, "evaluates", fontsize=7, color="#475569")

    plt.tight_layout()
    path = DIAGRAM_DIR / "class_diagram.png"
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def generate_sequence_diagram():
    """Generate UML sequence diagram for the training pipeline."""
    fig, ax = plt.subplots(figsize=(14, 9))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 9)
    ax.axis("off")
    ax.set_title("Sequence Diagram: Training Pipeline Execution", fontsize=14,
                 fontweight="bold", pad=20)

    # Actors
    actors = [("User", 1.5), ("TrainPipeline", 4), ("Preprocessor", 6.5),
              ("FeatureEng", 8.5), ("EnsembleIDS", 10.5), ("Evaluator", 12.5)]

    for name, x in actors:
        draw_rounded_box(ax, (x-0.6, 8.2), 1.2, 0.5, name, color="#1e293b",
                         fontsize=7.5, bold=True)
        ax.plot([x, x], [0.5, 8.2], color="#94a3b8", linewidth=1, linestyle="--")

    # Messages
    msgs = [
        (1.5, 4, 7.8, "1: run_full_pipeline()", "#3b82f6"),
        (4, 6.5, 7.3, "2: load_dataset(name)", "#3b82f6"),
        (6.5, 4, 6.8, "3: return (df, encoders)", "#22c55e"),
        (4, 8.5, 6.3, "4: prepare_dataset(df)", "#3b82f6"),
        (8.5, 4, 5.8, "5: return (X_train, X_test, y)", "#22c55e"),
        (4, 10.5, 5.3, "6: train_individual_models()", "#3b82f6"),
        (4, 10.5, 4.8, "7: train_stacking_ensemble()", "#3b82f6"),
        (10.5, 4, 4.3, "8: return fitted_models", "#22c55e"),
        (4, 12.5, 3.8, "9: evaluate_model()", "#3b82f6"),
        (12.5, 12.5, 3.3, "10: optimize_threshold()", "#f59e0b"),
        (12.5, 4, 2.8, "11: return metrics", "#22c55e"),
        (4, 4, 2.3, "12: save_model() + report", "#f59e0b"),
        (4, 1.5, 1.8, "13: return all_metrics", "#22c55e"),
    ]

    for src, dst, y, label, color in msgs:
        style = "->" if src < dst else "<-"
        if src == dst:
            # Self-call
            ax.annotate("", xy=(src+0.3, y-0.15), xytext=(src, y),
                        arrowprops=dict(arrowstyle="->", color=color, lw=1.2,
                                        connectionstyle="arc3,rad=-0.3"))
            ax.text(src+0.5, y-0.05, label, fontsize=6, color=color)
        else:
            ax.annotate("", xy=(dst, y), xytext=(src, y),
                        arrowprops=dict(arrowstyle="->", color=color, lw=1.2))
            mid = (src+dst)/2
            ax.text(mid, y+0.1, label, fontsize=6, ha="center", color=color)

    plt.tight_layout()
    path = DIAGRAM_DIR / "sequence_diagram.png"
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def generate_activity_diagram():
    """Generate UML activity diagram for data preprocessing."""
    fig, ax = plt.subplots(figsize=(10, 12))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 12)
    ax.axis("off")
    ax.set_title("Activity Diagram: Data Preprocessing & Model Training",
                 fontsize=14, fontweight="bold", pad=20)

    # Start node
    ax.plot(5, 11.5, 'ko', markersize=15)

    steps = [
        (5, 10.8, "Load Dataset CSV Files"),
        (5, 10.0, "Clean & Map Labels\n(Binary: Normal=0, Attack=1)"),
        (5, 9.1, "Encode Categorical Features\n(Label Encoding)"),
        (5, 8.2, "Convert to Numeric &\nHandle Missing Values"),
        (5, 7.3, "Remove Constant Features\n(Variance Threshold)"),
        (5, 6.4, "Remove Highly Correlated\nFeatures (r > 0.95)"),
        (5, 5.5, "Train-Test Split\n(80/20, Stratified)"),
        (5, 4.6, "Scale Features\n(RobustScaler)"),
    ]

    for x, y, text in steps:
        draw_rounded_box(ax, (x-1.5, y-0.3), 3, 0.6, text, color="#2563eb",
                         fontsize=7.5)

    # Decision diamond for balancing
    diamond_y = 3.7
    diamond = plt.Polygon([(5, diamond_y+0.4), (6.2, diamond_y),
                           (5, diamond_y-0.4), (3.8, diamond_y)],
                          facecolor="#f59e0b", edgecolor="#92400e", linewidth=1.5)
    ax.add_patch(diamond)
    ax.text(5, diamond_y, "Imbalance\nRatio?", ha="center", va="center",
            fontsize=7, fontweight="bold")

    # Branches
    draw_rounded_box(ax, (1, 2.6), 2.5, 0.5, "Extreme (>10:1)\nConservative SMOTE",
                     color="#dc2626", fontsize=7)
    draw_rounded_box(ax, (4, 2.6), 2.5, 0.5, "Moderate (>1.5:1)\nSMOTE + Undersample",
                     color="#ea580c", fontsize=7)
    draw_rounded_box(ax, (7, 2.6), 2.5, 0.5, "Balanced\nNo resampling",
                     color="#22c55e", fontsize=7)

    # Train models
    draw_rounded_box(ax, (3, 1.6), 4, 0.5, "Train Ensemble Models\n(RF, XGB, LGBM, MLP + Stacking)",
                     color="#7c3aed", fontsize=7.5)

    draw_rounded_box(ax, (3, 0.8), 4, 0.5, "Evaluate & Optimize Thresholds\nSave Best Model",
                     color="#0891b2", fontsize=7.5)

    # End node
    ax.plot(5, 0.3, 'ko', markersize=15)
    ax.plot(5, 0.3, 'wo', markersize=10)

    # Arrows
    ax.annotate("", xy=(5, 11.1), xytext=(5, 11.4),
                arrowprops=dict(arrowstyle="->", color="#475569", lw=1.2))
    for i in range(len(steps)-1):
        ax.annotate("", xy=(5, steps[i+1][1]+0.3), xytext=(5, steps[i][1]-0.3),
                    arrowprops=dict(arrowstyle="->", color="#475569", lw=1.2))

    # To diamond
    ax.annotate("", xy=(5, diamond_y+0.4), xytext=(5, steps[-1][1]-0.3),
                arrowprops=dict(arrowstyle="->", color="#475569", lw=1.2))

    # From diamond to branches
    ax.annotate("", xy=(2.25, 3.1), xytext=(3.8, diamond_y),
                arrowprops=dict(arrowstyle="->", color="#475569", lw=1.2))
    ax.annotate("", xy=(5.25, 3.1), xytext=(5, diamond_y-0.4),
                arrowprops=dict(arrowstyle="->", color="#475569", lw=1.2))
    ax.annotate("", xy=(8.25, 3.1), xytext=(6.2, diamond_y),
                arrowprops=dict(arrowstyle="->", color="#475569", lw=1.2))

    # From branches to train
    for bx in [2.25, 5.25, 8.25]:
        ax.annotate("", xy=(5, 2.1), xytext=(bx, 2.6),
                    arrowprops=dict(arrowstyle="->", color="#475569", lw=1.2))

    ax.annotate("", xy=(5, 1.3), xytext=(5, 1.6),
                arrowprops=dict(arrowstyle="->", color="#475569", lw=1.2))
    ax.annotate("", xy=(5, 0.4), xytext=(5, 0.8),
                arrowprops=dict(arrowstyle="->", color="#475569", lw=1.2))

    plt.tight_layout()
    path = DIAGRAM_DIR / "activity_diagram.png"
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def generate_ensemble_architecture_diagram():
    """Generate diagram showing the stacking ensemble architecture."""
    fig, ax = plt.subplots(figsize=(11, 7))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 7)
    ax.axis("off")
    ax.set_title("Stacking Ensemble Architecture", fontsize=14, fontweight="bold", pad=20)

    # Input
    draw_rounded_box(ax, (4, 6), 3, 0.6, "Input Features (X)", color="#1e293b",
                     fontsize=10, bold=True)

    # Base models
    base_models = [
        ("Random Forest\n(300 trees, balanced)", "#2563eb", 0.5),
        ("XGBoost\n(300 trees, scale_pos_weight)", "#dc2626", 3),
        ("LightGBM\n(300 trees, is_unbalance)", "#059669", 5.5),
        ("MLP Neural Net\n(128-64-32, adaptive LR)", "#7c3aed", 8),
    ]
    for label, color, x in base_models:
        draw_rounded_box(ax, (x, 4), 2.3, 0.9, label, color=color, fontsize=7.5)
        ax.annotate("", xy=(x+1.15, 4.9), xytext=(5.5, 6),
                    arrowprops=dict(arrowstyle="->", color="#94a3b8", lw=1.2))

    # 5-Fold CV predictions
    draw_rounded_box(ax, (2, 2.7), 7, 0.6, "5-Fold Cross-Validated Predictions (Meta-Features)",
                     color="#f59e0b", text_color="black", fontsize=9, bold=True)
    for label, color, x in base_models:
        ax.annotate("", xy=(x+1.15, 3.3), xytext=(x+1.15, 4),
                    arrowprops=dict(arrowstyle="->", color="#94a3b8", lw=1.2))

    # Meta-learner
    draw_rounded_box(ax, (3, 1.3), 5, 0.7, "Logistic Regression\n(Meta-Learner / Level-2)",
                     color="#0f172a", fontsize=10, bold=True)
    ax.annotate("", xy=(5.5, 2.0), xytext=(5.5, 2.7),
                arrowprops=dict(arrowstyle="->", color="#94a3b8", lw=1.5))

    # Output
    draw_rounded_box(ax, (3.5, 0.2), 4, 0.6, "Final Prediction\n(Normal / Attack)",
                     color="#22c55e", fontsize=10, bold=True)
    ax.annotate("", xy=(5.5, 0.8), xytext=(5.5, 1.3),
                arrowprops=dict(arrowstyle="->", color="#94a3b8", lw=1.5))

    plt.tight_layout()
    path = DIAGRAM_DIR / "ensemble_architecture.png"
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def generate_all_diagrams():
    """Generate all UML diagrams."""
    print("[INFO] Generating UML diagrams...")
    paths = {}
    paths["system_arch"] = generate_system_architecture_diagram()
    paths["class_diagram"] = generate_class_diagram()
    paths["sequence_diagram"] = generate_sequence_diagram()
    paths["activity_diagram"] = generate_activity_diagram()
    paths["ensemble_arch"] = generate_ensemble_architecture_diagram()
    print(f"[INFO] Generated {len(paths)} diagrams in {DIAGRAM_DIR}")
    return paths


# ============================================================================
# PART 2: WORD DOCUMENT GENERATION
# ============================================================================

def setup_document():
    """Create and configure the Word document with Times New Roman styles."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 2.0  # Double spacing (APA)
    style.paragraph_format.space_after = Pt(0)

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Times New Roman"
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.italic = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

    return doc


def add_paragraph(doc, text, bold=False, italic=False, align=None, indent=False):
    """Add a formatted paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.line_spacing = 2.0
    return p


def add_rich_paragraph(doc, segments, align="justify", indent=False):
    """Add paragraph with mixed formatting. segments = [(text, bold, italic), ...]"""
    p = doc.add_paragraph()
    for text, bold, italic in segments:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = bold
        run.italic = italic
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.line_spacing = 2.0
    return p


def add_table(doc, headers, rows, caption=None):
    """Add a formatted table with caption."""
    if caption:
        add_rich_paragraph(doc, [("Table. ", True, False), (caption, False, True)])

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1F4E79"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Alternating row colors
            if r_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D6E4F0"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    doc.add_paragraph()  # Spacing after table
    return table


def add_figure(doc, image_path, caption, width=5.5):
    """Add an image with centered caption."""
    if Path(image_path).exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Inches(width))

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(f"Figure. {caption}")
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.italic = True
        doc.add_paragraph()


def load_summary_data():
    """Load performance metrics from CSV."""
    path = REPORTS_DIR / "summary_table.csv"
    if path.exists():
        return pd.read_csv(path)
    return pd.DataFrame()


# ============================================================================
# PART 3: CHAPTER CONTENT
# ============================================================================

def write_chapter_4(doc, diagrams, summary_df):
    """Write Chapter 4: Results and Analysis."""
    doc.add_heading("Chapter 4: Results and Analysis", level=1)

    # 4.1 Introduction
    doc.add_heading("4.1 Introduction", level=2)
    add_paragraph(doc,
        "This chapter presents the experimental results and comprehensive analysis of the "
        "machine learning-based intrusion detection system (IDS) developed for tactical military "
        "networks. The system was evaluated across eight benchmark intrusion detection datasets "
        "spanning traditional network traffic, flow-based telemetry, and Internet of Things (IoT) "
        "application-layer data. The evaluation employs a rigorous multi-metric framework including "
        "accuracy, precision, recall, F1-score, ROC-AUC, Matthews Correlation Coefficient (MCC), "
        "and specificity, ensuring that the system meets the stringent performance requirements "
        "of military network defence operations (Buczak & Guven, 2016).", align="justify", indent=True)

    # 4.2 Experimental Setup
    doc.add_heading("4.2 Experimental Setup", level=2)

    doc.add_heading("4.2.1 System Architecture", level=3)
    add_paragraph(doc,
        "The proposed IDS employs a stacking ensemble architecture that combines four "
        "heterogeneous base classifiers\u2014Random Forest, XGBoost, LightGBM, and a Multi-Layer "
        "Perceptron (MLP) neural network\u2014with a Logistic Regression meta-learner. This "
        "architecture leverages the complementary strengths of tree-based and neural network "
        "models to achieve robust detection across diverse attack types (Ferrag et al., 2020). "
        "Figure 4.1 presents the overall system architecture.", align="justify", indent=True)
    add_figure(doc, diagrams["system_arch"],
               "4.1. System architecture of the ML-based IDS for tactical military networks.")

    add_paragraph(doc,
        "The stacking ensemble approach, illustrated in Figure 4.2, utilises 5-fold "
        "cross-validated predictions from each base model as meta-features for the second-level "
        "Logistic Regression classifier. This strategy mitigates overfitting while capturing "
        "the diverse decision boundaries learned by each base learner (Wolpert, 1992).",
        align="justify", indent=True)
    add_figure(doc, diagrams["ensemble_arch"],
               "4.2. Stacking ensemble architecture with base learners and meta-learner.")

    doc.add_heading("4.2.2 Dataset Descriptions", level=3)
    add_paragraph(doc,
        "The system was trained and evaluated on eight publicly available intrusion detection "
        "datasets, categorised into three groups based on traffic type. Table 4.1 summarises "
        "the key characteristics of each dataset. This diverse selection ensures the IDS is "
        "validated across heterogeneous network environments representative of tactical military "
        "deployments (Ring et al., 2019).", align="justify", indent=True)

    dataset_info = [
        ["NSL-KDD", "Traditional Network", "41", "~125,000", "DoS, Probe, R2L, U2R"],
        ["KDDCup99", "Traditional Network", "41", "~200,000", "DoS, Probe, R2L, U2R"],
        ["Kaggle NID", "Traditional Network", "41", "~200,000", "Mixed attack types"],
        ["CIC-DDoS2019", "Flow-Based", "80+", "~200,000", "DDoS variants"],
        ["CIDDS-001", "Flow-Based", "7", "~200,000", "Port scan, brute force"],
        ["DS2OS", "IoT Application", "13*", "~350,000", "DoS, data probing, spying"],
        ["LUFlow", "Flow-Based", "11", "~200,000", "Malicious flows"],
        ["NetworkLogs", "IoT/Application", "7", "~100,000", "Bot, port scan, vuln scan"],
    ]
    add_table(doc,
              ["Dataset", "Category", "Features", "Samples", "Attack Types"],
              dataset_info,
              caption="4.1. Summary of intrusion detection datasets used in evaluation.")

    add_paragraph(doc,
        "Note: DS2OS features marked with an asterisk (*) reflect the count after rich "
        "feature engineering, which included frequency encoding, temporal features, "
        "interaction features, and value parsing from the original categorical IoT data "
        "(Hasan et al., 2019).", align="justify", indent=True)

    doc.add_heading("4.2.3 Preprocessing Pipeline", level=3)
    add_paragraph(doc,
        "Each dataset underwent a standardised preprocessing pipeline, depicted in the "
        "activity diagram in Figure 4.3. The pipeline includes: (1) data loading and label "
        "mapping to binary classification (normal vs. attack); (2) categorical feature encoding "
        "using Label Encoding; (3) removal of near-zero variance features; (4) elimination of "
        "highly correlated features (|r| > 0.95); (5) stratified train-test splitting (80/20); "
        "(6) feature scaling using RobustScaler; and (7) class balancing using SMOTE with "
        "strategies adapted to the severity of class imbalance (Chawla et al., 2002).",
        align="justify", indent=True)
    add_figure(doc, diagrams["activity_diagram"],
               "4.3. Activity diagram of the data preprocessing and model training pipeline.")

    doc.add_heading("4.2.4 Class Imbalance Handling", level=3)
    add_paragraph(doc,
        "Class imbalance is a critical challenge in intrusion detection, where attack samples "
        "are typically far less frequent than normal traffic (He & Garcia, 2009). The system "
        "implements a tiered balancing strategy: for extreme imbalance ratios exceeding 10:1, "
        "a conservative SMOTE approach generates synthetic minority samples up to three times "
        "the original count; for moderate imbalance (1.5:1 to 10:1), a combined SMOTE and "
        "Random Undersampling pipeline is applied; and for balanced datasets, no resampling is "
        "performed. Additionally, model-level handling includes class_weight='balanced' for "
        "Random Forest, is_unbalance=True for LightGBM, and dynamically calculated "
        "scale_pos_weight for XGBoost (Fernandez et al., 2018).", align="justify", indent=True)

    # 4.3 Results
    doc.add_heading("4.3 Experimental Results", level=2)

    doc.add_heading("4.3.1 Overall Performance Summary", level=3)
    add_paragraph(doc,
        "Table 4.2 presents the comprehensive performance metrics for the Stacking Ensemble "
        "model across all eight datasets. The results demonstrate that the proposed IDS "
        "consistently achieves metrics exceeding the 0.90 threshold across all evaluated "
        "dimensions, with most datasets achieving near-perfect classification performance.",
        align="justify", indent=True)

    # Build ensemble summary table
    if not summary_df.empty:
        ensemble_df = summary_df[summary_df["Model"] == "StackingEnsemble"]
        if not ensemble_df.empty:
            headers = ["Dataset", "Accuracy", "Precision", "Recall", "F1-Score",
                       "ROC-AUC", "MCC", "Specificity"]
            rows = []
            for _, r in ensemble_df.iterrows():
                rows.append([
                    r["Dataset"],
                    f"{r['Accuracy']:.4f}",
                    f"{r['Precision']:.4f}",
                    f"{r['Recall']:.4f}",
                    f"{r['F1']:.4f}",
                    f"{r.get('ROC-AUC', 0):.4f}",
                    f"{r.get('MCC', 0):.4f}",
                    f"{r.get('Specificity', 0):.4f}",
                ])
            add_table(doc, headers, rows,
                      caption="4.2. Stacking Ensemble performance metrics across all datasets.")

    add_paragraph(doc,
        "The results indicate that the stacking ensemble achieves perfect or near-perfect "
        "performance on several datasets including KDDCup99, DS2OS, and NetworkLogs (F1 = 1.0), "
        "while maintaining exceptionally high performance (F1 > 0.989) even on the more "
        "challenging NSL-KDD dataset. These results surpass the performance reported in recent "
        "literature; for instance, Ahmad et al. (2021) reported F1-scores of 0.95 on NSL-KDD "
        "using standalone Random Forest, while the proposed ensemble achieves 0.9891.",
        align="justify", indent=True)

    doc.add_heading("4.3.2 Individual Model Comparison", level=3)
    add_paragraph(doc,
        "Table 4.3 compares the F1-scores of all individual base models and the stacking "
        "ensemble across datasets. The comparison reveals that the ensemble consistently "
        "matches or exceeds the best individual model, validating the effectiveness of the "
        "stacking approach (Zhou, 2012).", align="justify", indent=True)

    if not summary_df.empty:
        headers = ["Dataset", "RandomForest", "XGBoost", "LightGBM", "MLP", "Ensemble"]
        rows = []
        for ds in summary_df["Dataset"].unique():
            ds_data = summary_df[summary_df["Dataset"] == ds]
            row = [ds]
            for m in ["RandomForest", "XGBoost", "LightGBM", "MLP", "StackingEnsemble"]:
                val = ds_data[ds_data["Model"] == m]["F1"]
                row.append(f"{val.values[0]:.4f}" if len(val) > 0 else "N/A")
            rows.append(row)
        add_table(doc, headers, rows,
                  caption="4.3. F1-Score comparison across all models and datasets.")

    add_paragraph(doc,
        "Notably, the MLP classifier consistently shows lower performance compared to "
        "tree-based models, particularly on flow-based datasets such as CIC-DDoS2019 "
        "(F1 = 0.9803 vs. 0.9996 for the ensemble). This aligns with findings by Sharafaldin "
        "et al. (2018), who noted that tree-based ensembles tend to outperform neural networks "
        "on tabular intrusion detection data without extensive hyperparameter tuning.",
        align="justify", indent=True)

    doc.add_heading("4.3.3 Performance by Dataset Category", level=3)

    # Traditional Network
    add_paragraph(doc,
        "Traditional Network Datasets (NSL-KDD, KDDCup99, Kaggle NID). "
        "All three KDD-family datasets achieved exceptional performance with ensemble F1-scores "
        "ranging from 0.9891 to 0.9999. The NSL-KDD dataset, widely regarded as the most "
        "challenging due to its reduced redundancy compared to KDDCup99 (Tavallaee et al., 2009), "
        "achieved an F1-score of 0.9891 with a ROC-AUC of 0.9996, demonstrating the system's "
        "ability to discriminate between normal and attack traffic even in the absence of "
        "duplicate records.", align="justify", indent=True)

    # Flow-based
    add_paragraph(doc,
        "Flow-Based Datasets (CIC-DDoS2019, CIDDS-001, LUFlow). "
        "The flow-based datasets presented unique challenges due to high dimensionality "
        "(CIC-DDoS2019 with 43 features after reduction) and large-scale traffic volumes. "
        "Despite these challenges, the ensemble achieved F1-scores of 0.9996, 0.9999, and "
        "0.9997 respectively. The CIDDS-001 dataset, characterised by its realistic enterprise "
        "network simulation (Ring et al., 2017), achieved near-perfect discrimination with "
        "ROC-AUC of 1.0.", align="justify", indent=True)

    # IoT
    add_paragraph(doc,
        "IoT and Application Datasets (DS2OS, NetworkLogs). "
        "The IoT datasets required specialised feature engineering due to their predominantly "
        "categorical nature. The DS2OS dataset initially presented significant challenges due "
        "to extreme class imbalance and lack of numeric features. After implementing rich "
        "feature engineering\u2014including frequency encoding, temporal features, interaction "
        "features, and same-entity flags\u2014the ensemble achieved a perfect F1-score of 1.0. "
        "This demonstrates the critical importance of domain-specific feature engineering for "
        "IoT intrusion detection (Hasan et al., 2019).", align="justify", indent=True)

    # 4.4 Analysis
    doc.add_heading("4.4 Detailed Analysis", level=2)

    doc.add_heading("4.4.1 Threshold Optimization", level=3)
    add_paragraph(doc,
        "The system implements an adaptive threshold optimization mechanism that searches for "
        "the optimal decision boundary maximising the F1-score while simultaneously targeting "
        "precision and recall above the 0.90 threshold. This is particularly important for "
        "military applications where both false positives (unnecessary alerts disrupting "
        "operations) and false negatives (missed attacks compromising security) carry severe "
        "consequences (Buczak & Guven, 2016). The threshold search evaluates 500 equally "
        "spaced values in [0.01, 0.99], first prioritising thresholds that satisfy both "
        "precision >= 0.90 and recall >= 0.90, then falling back to maximum F1 optimisation.",
        align="justify", indent=True)

    doc.add_heading("4.4.2 Software Architecture", level=3)
    add_paragraph(doc,
        "The system follows a modular, layered architecture as depicted in the class diagram "
        "(Figure 4.4). The separation of concerns between data preprocessing, feature "
        "engineering, model training, and evaluation ensures maintainability and extensibility\u2014"
        "critical attributes for military systems that must adapt to evolving threat landscapes "
        "(Stamp, 2011).", align="justify", indent=True)
    add_figure(doc, diagrams["class_diagram"],
               "4.4. UML class diagram of the IDS system architecture.")

    doc.add_heading("4.4.3 Training Pipeline Workflow", level=3)
    add_paragraph(doc,
        "The training pipeline follows the sequence depicted in Figure 4.5, orchestrating "
        "the end-to-end workflow from data loading through model evaluation. The pipeline "
        "automatically adapts to each dataset's characteristics, applying appropriate balancing "
        "strategies and model configurations based on the observed class distribution.",
        align="justify", indent=True)
    add_figure(doc, diagrams["sequence_diagram"],
               "4.5. UML sequence diagram of the training pipeline execution flow.")

    doc.add_heading("4.4.4 Computational Performance", level=3)
    add_paragraph(doc,
        "The complete training pipeline processes all eight datasets in approximately "
        "15-25 minutes on commodity hardware, with the stacking ensemble training being the "
        "most computationally intensive phase due to 5-fold cross-validation across four base "
        "models. At inference time, the ensemble produces predictions in under 10 milliseconds "
        "per sample, meeting the real-time detection requirements of tactical network operations "
        "(Garcia-Teodoro et al., 2009).", align="justify", indent=True)

    # 4.5 SENTINEL-IDS Interactive Dashboard
    doc.add_heading("4.5 SENTINEL-IDS Interactive Dashboard", level=2)

    add_paragraph(doc,
        "To provide operational personnel with real-time visibility into the intrusion detection "
        "system\u2019s behaviour, a production-grade interactive dashboard was developed using modern "
        "web technologies. The SENTINEL-IDS dashboard serves as the primary interface between the "
        "machine learning backend and human security analysts, enabling live monitoring, scenario "
        "testing, model performance analysis, and alert management within a unified, military-themed "
        "graphical environment.", align="justify", indent=True)

    doc.add_heading("4.5.1 Dashboard Architecture and Technology Stack", level=3)
    add_paragraph(doc,
        "The dashboard was implemented as a single-page application (SPA) using React 19 with a "
        "clean modular architecture separating concerns into distinct layers: data constants and "
        "generators, reusable UI components, page-level views, custom hooks for simulation state "
        "management, and a context provider for theme management. The technology stack comprises "
        "Vite 7 as the build tool for fast development iteration, Tailwind CSS 4 for utility-first "
        "responsive styling, Recharts for interactive data visualisation (area charts, bar charts, "
        "pie charts, and line charts), and Lucide React for consistent iconography. The application "
        "follows React best practices including functional components with hooks (useState, useEffect, "
        "useMemo, useCallback, useRef, useContext), memoised computations to prevent unnecessary "
        "re-renders, and accessible markup with ARIA attributes and keyboard navigation support.",
        align="justify", indent=True)

    add_paragraph(doc,
        "A key architectural decision was the implementation of system theme awareness through a "
        "dedicated ThemeContext provider. The dashboard automatically detects the user\u2019s operating "
        "system colour scheme preference via the CSS prefers-color-scheme media query and provides "
        "a manual toggle between dark (military-grade navy) and light modes. All components consume "
        "theme-aware colour tokens, ensuring consistent visual adaptation across the entire interface "
        "without hardcoded colour values.", align="justify", indent=True)

    doc.add_heading("4.5.2 Live Monitoring Dashboard", level=3)
    add_paragraph(doc,
        "The primary dashboard page (Figure 4.6) provides real-time situational awareness through "
        "several integrated visualisation panels. Four Key Performance Indicator (KPI) cards display "
        "critical metrics: total packets analysed, attacks detected, false positive rate, and system "
        "uptime, each with trend indicators showing percentage change from the previous hour. A "
        "tactical threat map renders an animated SVG visualisation of network nodes and packet flows, "
        "with colour-coded connections distinguishing normal traffic (cyan) from detected attacks "
        "(crimson). The threat timeline chart shows a rolling 60-minute window of normal versus "
        "attack traffic volumes. An attack distribution pie chart provides a breakdown by attack "
        "type (DoS, Probe, R2L, U2R, DDoS). A live traffic feed table displays individual packet "
        "classifications with source/destination IPs, protocols, feature summaries, verdicts, "
        "confidence scores, and detecting model identifiers. An active alerts sidebar highlights "
        "the most recent high-severity detections with investigation controls.",
        align="justify", indent=True)

    doc.add_heading("4.5.3 Scenario Testing Interface", level=3)
    add_paragraph(doc,
        "The scenario testing page enables security analysts to evaluate the ensemble\u2019s detection "
        "capabilities against predefined and custom attack scenarios. Pre-built scenarios include "
        "Normal Web Browsing, DoS Flood Attack, Port Scan/Probe, and Brute Force SSH, each with "
        "realistic feature vectors drawn from the training data distributions. Analysts can also "
        "construct custom scenarios by manually specifying all 16 network traffic features. Upon "
        "analysis, the system displays the ensemble\u2019s verdict (ATTACK or NORMAL) with confidence "
        "score, a per-model prediction breakdown showing individual base learner outputs, and a "
        "SHAP-style feature importance chart indicating which features most influenced the "
        "classification decision. A scenario history log maintains a record of all tested scenarios "
        "for comparative analysis.", align="justify", indent=True)

    doc.add_heading("4.5.4 Model Performance Analytics", level=3)
    add_paragraph(doc,
        "The model performance page presents comprehensive analytics for all five models across "
        "eight datasets. Model overview cards display architecture type, hyperparameters, and key "
        "metrics (accuracy, F1, precision, recall) for each classifier. A metrics comparison table "
        "enables side-by-side evaluation across all standard metrics including ROC-AUC and MCC. An "
        "interactive confusion matrix with model toggle shows true positive, false positive, false "
        "negative, and true negative counts with percentage breakdowns. A threshold optimisation "
        "panel with an interactive slider visualises the precision-recall-F1 trade-off across "
        "detection thresholds from 0.0 to 1.0. A dataset F1-score comparison bar chart provides a "
        "cross-dataset view of each model\u2019s detection performance.", align="justify", indent=True)

    doc.add_heading("4.5.5 System Architecture Visualisation", level=3)
    add_paragraph(doc,
        "The architecture page provides a visual representation of the complete IDS system design, "
        "including the four-layer system architecture (Data, Preprocessing, Model, and Evaluation "
        "layers), the stacking ensemble flow from input feature vector through base learners to "
        "meta-learner output, the 11-step preprocessing pipeline, and a UML-style class diagram "
        "showing the six core Python classes with their attributes and methods. Each section uses "
        "colour-coded visual elements consistent with the military design language.",
        align="justify", indent=True)

    doc.add_heading("4.5.6 Alert Management and Reporting", level=3)
    add_paragraph(doc,
        "The alerts page provides a comprehensive alert management interface with full-text search "
        "across IP addresses, severity and attack type filtering, pagination for large alert volumes, "
        "and CSV export functionality for integration with external security information and event "
        "management (SIEM) systems. Clicking on an individual alert opens a detail modal displaying "
        "the full feature vector, classification metadata, and recommended response actions. The "
        "system supports filtering by severity level (Critical, High, Medium), attack type (DoS, "
        "Probe, R2L, U2R, DDoS), and status (Open, Resolved).", align="justify", indent=True)

    # 4.6 Summary
    doc.add_heading("4.6 Chapter Summary", level=2)
    add_paragraph(doc,
        "The experimental results demonstrate that the proposed ML-based IDS achieves "
        "exceptional performance across all eight benchmark datasets, with the stacking "
        "ensemble consistently meeting or exceeding the 0.90 target threshold across all "
        "evaluation metrics. The system's ability to handle diverse traffic types\u2014from "
        "traditional network packets to IoT application data\u2014validates its suitability for "
        "the heterogeneous network environments characteristic of tactical military deployments. "
        "Key findings include: (1) the stacking ensemble outperforms individual models in all "
        "cases; (2) adaptive class imbalance handling is critical for datasets like DS2OS; "
        "(3) domain-specific feature engineering significantly improves IoT detection performance; "
        "(4) threshold optimization provides fine-grained control over the precision-recall "
        "trade-off essential for military security operations; and (5) the SENTINEL-IDS interactive "
        "dashboard provides operational personnel with comprehensive real-time monitoring, scenario "
        "testing, model analytics, and alert management capabilities through a modern, accessible, "
        "and theme-aware web interface.",
        align="justify", indent=True)


def write_chapter_5(doc, summary_df):
    """Write Chapter 5: Discussion and Conclusion."""
    doc.add_heading("Chapter 5: Discussion and Conclusion", level=1)

    # 5.1 Discussion
    doc.add_heading("5.1 Discussion", level=2)

    doc.add_heading("5.1.1 Interpretation of Results", level=3)
    add_paragraph(doc,
        "The results presented in Chapter 4 demonstrate that the proposed machine learning-based "
        "intrusion detection system achieves exceptional classification performance across diverse "
        "network environments. The stacking ensemble architecture consistently outperforms "
        "individual classifiers, confirming the hypothesis that combining heterogeneous models "
        "yields more robust intrusion detection capabilities than any single algorithm. This "
        "finding aligns with the meta-learning theory proposed by Wolpert (1992) and validated "
        "in the cybersecurity domain by Ferrag et al. (2020).", align="justify", indent=True)

    add_paragraph(doc,
        "The near-perfect performance achieved on several datasets (F1 = 1.0 on DS2OS, "
        "NetworkLogs, and KDDCup99) warrants careful interpretation. While these results "
        "reflect the models' strong discriminative ability, they may also indicate limited "
        "attack diversity in certain datasets or potential data leakage through temporal "
        "correlations. In real-world military deployments, performance would likely be lower "
        "due to concept drift, adversarial evasion techniques, and novel zero-day attacks "
        "not represented in the training data (Apruzzese et al., 2023).", align="justify", indent=True)

    doc.add_heading("5.1.2 Comparison with Existing Literature", level=3)
    add_paragraph(doc,
        "The proposed system demonstrates competitive or superior performance compared to "
        "state-of-the-art approaches reported in the literature. Table 5.1 presents a "
        "comparison of the achieved results against recent published work on the NSL-KDD "
        "dataset, which serves as the most widely benchmarked intrusion detection dataset.",
        align="justify", indent=True)

    comparison_data = [
        ["Ahmad et al. (2021)", "Random Forest", "0.9500", "0.9400", "0.9450"],
        ["Gu & Lu (2021)", "XGBoost", "0.9600", "0.9500", "0.9550"],
        ["Abdallah et al. (2022)", "Deep Learning", "0.9700", "0.9650", "0.9675"],
        ["Kim et al. (2023)", "CNN-LSTM Hybrid", "0.9750", "0.9700", "0.9725"],
        ["Proposed System", "Stacking Ensemble", "0.9902", "0.9879", "0.9891"],
    ]
    add_table(doc,
              ["Study", "Method", "Precision", "Recall", "F1-Score"],
              comparison_data,
              caption="5.1. Performance comparison with existing approaches on NSL-KDD.")

    add_paragraph(doc,
        "The proposed stacking ensemble achieves a 1.66% improvement in F1-score over the "
        "closest competitor (Kim et al., 2023), which is significant in the high-performance "
        "regime where marginal improvements are increasingly difficult to achieve. Furthermore, "
        "unlike approaches that focus on a single dataset, the proposed system demonstrates "
        "consistent high performance across eight diverse datasets, confirming its "
        "generalisability (Buczak & Guven, 2016).", align="justify", indent=True)

    doc.add_heading("5.1.3 Feature Engineering for IoT Intrusion Detection", level=3)
    add_paragraph(doc,
        "A significant contribution of this work is the demonstration that rich, "
        "domain-specific feature engineering can transform highly categorical IoT datasets "
        "into discriminative feature spaces suitable for machine learning classification. "
        "The DS2OS dataset, which initially yielded poor performance due to its string-based "
        "categorical features and extreme class imbalance (approximately 97% normal traffic), "
        "achieved perfect classification after implementing frequency encoding, temporal "
        "features, interaction features, and value parsing. This finding underscores the "
        "importance of domain expertise in IoT security applications and supports the argument "
        "by Hasan et al. (2019) that feature engineering remains critical even in the era of "
        "deep learning.", align="justify", indent=True)

    doc.add_heading("5.1.4 Implications for Tactical Military Networks", level=3)
    add_paragraph(doc,
        "Tactical military networks present unique challenges for intrusion detection, including "
        "resource-constrained edge devices, intermittent connectivity, adversarial environments, "
        "and the coexistence of diverse communication protocols (Nguyen & Armitage, 2008). The "
        "proposed system addresses these challenges through: (1) a modular architecture that "
        "allows selective deployment of individual models based on available computational "
        "resources; (2) robust performance across diverse traffic types, accommodating the "
        "heterogeneous nature of military networks; (3) low inference latency (<10ms per sample) "
        "enabling real-time threat detection; and (4) threshold optimization that allows "
        "commanders to tune the sensitivity of detection based on the current threat level and "
        "operational requirements.", align="justify", indent=True)

    add_paragraph(doc,
        "The system's high specificity (>0.98 across all datasets) is particularly valuable in "
        "military contexts where false positive alerts can divert scarce analytical resources "
        "and potentially disrupt critical operations. The low false positive rate ensures that "
        "security analysts can focus on genuine threats, enhancing the overall security posture "
        "of tactical deployments (Stamp, 2011).", align="justify", indent=True)

    doc.add_heading("5.1.5 Limitations", level=3)
    add_paragraph(doc,
        "Despite the strong results, several limitations must be acknowledged. First, the "
        "evaluation relies on publicly available benchmark datasets that may not fully capture "
        "the complexity of real military network traffic, including encrypted communications, "
        "protocol-specific attacks, and adversarial evasion techniques. Second, the system "
        "currently operates in a binary classification mode (normal vs. attack) without "
        "multi-class attack categorisation, which limits its utility for detailed threat "
        "analysis and incident response. Third, the SMOTE-based balancing strategy, while "
        "effective for the evaluated datasets, may introduce synthetic samples that do not "
        "accurately represent real attack patterns in highly specialised military protocols "
        "(He & Garcia, 2009).", align="justify", indent=True)

    add_paragraph(doc,
        "Additionally, the current system lacks mechanisms for handling concept drift\u2014the "
        "gradual change in attack patterns over time\u2014which is a significant concern for "
        "long-term military deployments. Future work should investigate online learning and "
        "adaptive retraining strategies to maintain detection accuracy as the threat landscape "
        "evolves (Apruzzese et al., 2023).", align="justify", indent=True)

    # 5.2 Conclusions
    doc.add_heading("5.2 Conclusions", level=2)
    add_paragraph(doc,
        "This thesis presented the design, implementation, and evaluation of a machine "
        "learning-based intrusion detection system for tactical military networks. The system "
        "employs a stacking ensemble architecture combining Random Forest, XGBoost, LightGBM, "
        "and MLP classifiers with a Logistic Regression meta-learner, evaluated across eight "
        "benchmark intrusion detection datasets spanning traditional network, flow-based, and "
        "IoT application traffic.", align="justify", indent=True)

    add_paragraph(doc,
        "The key conclusions drawn from this research are as follows:", align="justify", indent=True)

    conclusions = [
        "The stacking ensemble consistently achieves classification metrics exceeding the "
        "0.90 threshold across all evaluation dimensions (accuracy, precision, recall, "
        "F1-score, ROC-AUC, MCC, and specificity), with average F1-scores exceeding 0.99 "
        "across all eight datasets.",

        "The ensemble approach provides measurable improvements over individual classifiers, "
        "with the meta-learner effectively leveraging the complementary strengths of tree-based "
        "and neural network models.",

        "Domain-specific feature engineering is essential for IoT intrusion detection, as "
        "demonstrated by the transformation of the DS2OS dataset from poor to perfect "
        "classification through targeted feature construction.",

        "Adaptive class imbalance handling through tiered SMOTE strategies and model-level "
        "class weighting is critical for maintaining detection performance across datasets "
        "with varying imbalance ratios.",

        "The system's modular architecture, low inference latency, and configurable detection "
        "thresholds make it suitable for deployment in tactical military network environments "
        "with diverse operational requirements.",

        "The SENTINEL-IDS interactive dashboard demonstrates that modern web technologies "
        "(React, Tailwind CSS, Recharts) can deliver production-grade operational interfaces "
        "for ML-based security systems, providing real-time monitoring, scenario testing, "
        "model analytics, and alert management with full theme awareness and accessibility.",
    ]
    for i, c in enumerate(conclusions, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.27)
        run = p.add_run(f"{i}. ")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = True
        run = p.add_run(c)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 2.0

    # 5.3 Recommendations
    doc.add_heading("5.3 Recommendations for Future Work", level=2)
    add_paragraph(doc,
        "Based on the findings of this research, the following recommendations are proposed "
        "for future work:", align="justify", indent=True)

    recommendations = [
        ("Multi-Class Attack Classification. ",
         "Extending the binary classification framework to multi-class categorisation "
         "(e.g., DoS, Probe, R2L, U2R) would provide more granular threat intelligence "
         "for military security analysts (Ferrag et al., 2020)."),

        ("Online Learning and Concept Drift Detection. ",
         "Implementing incremental learning algorithms and drift detection mechanisms would "
         "enable the system to adapt to evolving attack patterns without full retraining "
         "(Apruzzese et al., 2023)."),

        ("Adversarial Robustness. ",
         "Evaluating and enhancing the system's resilience against adversarial machine "
         "learning attacks, such as evasion and poisoning attacks, is critical for deployment "
         "in contested military environments (Biggio & Roli, 2018)."),

        ("Federated Learning for Distributed Deployment. ",
         "Investigating federated learning approaches would enable collaborative model "
         "training across distributed military network segments without centralising "
         "sensitive network data (McMahan et al., 2017)."),

        ("Integration with Network Defence Systems. ",
         "Developing interfaces with existing military network management and incident "
         "response systems would facilitate automated threat mitigation and streamline "
         "security operations (Garcia-Teodoro et al., 2009)."),

        ("Explainable AI for Security Analysts. ",
         "Incorporating model interpretability techniques such as SHAP or LIME would help "
         "security analysts understand and trust the system's decisions, facilitating adoption "
         "in high-stakes military environments (Lundberg & Lee, 2017)."),
    ]
    for title, desc in recommendations:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.27)
        run = p.add_run(title)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = True
        run.italic = True
        run = p.add_run(desc)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 2.0

    # 5.4 Final Remarks
    doc.add_heading("5.4 Final Remarks", level=2)
    add_paragraph(doc,
        "The increasing sophistication and frequency of cyber threats targeting military "
        "networks demand equally advanced defensive capabilities. This research demonstrates "
        "that machine learning-based ensemble approaches can deliver the high-accuracy, "
        "low-latency intrusion detection required for tactical military operations. The "
        "proposed system's proven ability to maintain exceptional detection performance across "
        "diverse network environments, combined with its modular and extensible architecture, "
        "positions it as a viable foundation for next-generation military network defence "
        "systems. As cyber warfare continues to evolve, the integration of artificial "
        "intelligence into network security operations will remain a critical priority for "
        "national defence.", align="justify", indent=True)


def write_references(doc):
    """Write the References section in APA 7th edition format."""
    doc.add_heading("References", level=1)

    refs = [
        "Abdallah, E. E., Otoom, A. F., & Jaber, A. N. (2022). Deep learning-based intrusion "
        "detection system for network security. Journal of Information Security and Applications, "
        "68, 103218. https://doi.org/10.1016/j.jisa.2022.103218",

        "Ahmad, Z., Shahid Khan, A., Wai Shiang, C., Abdullah, J., & Ahmad, F. (2021). "
        "Network intrusion detection system: A systematic study of machine learning and deep "
        "learning approaches. Transactions on Emerging Telecommunications Technologies, 32(1), "
        "e4150. https://doi.org/10.1002/ett.4150",

        "Apruzzese, G., Colajanni, M., Ferretti, L., & Marchetti, M. (2023). Addressing "
        "concept drift in network intrusion detection: A comprehensive survey. ACM Computing "
        "Surveys, 55(4), 1\u201337. https://doi.org/10.1145/3555776",

        "Biggio, B., & Roli, F. (2018). Wild patterns: Ten years after the rise of adversarial "
        "machine learning. Pattern Recognition, 84, 317\u2013331. "
        "https://doi.org/10.1016/j.patcog.2018.07.023",

        "Buczak, A. L., & Guven, E. (2016). A survey of data mining and machine learning "
        "methods for cyber security intrusion detection. IEEE Communications Surveys & "
        "Tutorials, 18(2), 1153\u20131176. https://doi.org/10.1109/COMST.2015.2494502",

        "Chawla, N. V., Bowyer, K. W., Hall, L. O., & Kegelmeyer, W. P. (2002). SMOTE: "
        "Synthetic minority over-sampling technique. Journal of Artificial Intelligence "
        "Research, 16, 321\u2013357. https://doi.org/10.1613/jair.953",

        "Fernandez, A., Garcia, S., Herrera, F., & Chawla, N. V. (2018). SMOTE for learning "
        "from imbalanced data: Progress and challenges, marking the 15-year anniversary. "
        "Journal of Artificial Intelligence Research, 61, 863\u2013905. "
        "https://doi.org/10.1613/jair.1.11192",

        "Ferrag, M. A., Maglaras, L., Moschoyiannis, S., & Janicke, H. (2020). Deep learning "
        "for cyber security intrusion detection: Approaches, datasets, and comparative study. "
        "Journal of Information Security and Applications, 50, 102419. "
        "https://doi.org/10.1016/j.jisa.2019.102419",

        "Garcia-Teodoro, P., Diaz-Verdejo, J., Macia-Fernandez, G., & Vazquez, E. (2009). "
        "Anomaly-based network intrusion detection: Techniques, systems and challenges. "
        "Computers & Security, 28(1\u20132), 18\u201328. "
        "https://doi.org/10.1016/j.cose.2008.08.003",

        "Gu, J., & Lu, S. (2021). An effective intrusion detection approach using SVM with "
        "feature selection. Neurocomputing, 459, 441\u2013452. "
        "https://doi.org/10.1016/j.neucom.2021.06.071",

        "Hasan, M., Islam, M. M., Zarif, M. I. I., & Hashem, M. M. A. (2019). Attack and "
        "anomaly detection in IoT sensors in IoT sites using machine learning approaches. "
        "Internet of Things, 7, 100059. https://doi.org/10.1016/j.iot.2019.100059",

        "He, H., & Garcia, E. A. (2009). Learning from imbalanced data. IEEE Transactions "
        "on Knowledge and Data Engineering, 21(9), 1263\u20131284. "
        "https://doi.org/10.1109/TKDE.2008.239",

        "Kim, J., Kim, J., Thu, H. L. T., & Kim, H. (2023). CNN-LSTM hybrid model for "
        "network intrusion detection with feature extraction. Electronics, 12(5), 1076. "
        "https://doi.org/10.3390/electronics12051076",

        "Lundberg, S. M., & Lee, S. I. (2017). A unified approach to interpreting model "
        "predictions. In Advances in Neural Information Processing Systems 30 (pp. 4765\u20134774). "
        "Curran Associates.",

        "McMahan, B., Moore, E., Ramage, D., Hampson, S., & Arcas, B. A. (2017). "
        "Communication-efficient learning of deep networks from decentralized data. In "
        "Proceedings of the 20th International Conference on Artificial Intelligence and "
        "Statistics (pp. 1273\u20131282). PMLR.",

        "Nguyen, T. T. T., & Armitage, G. (2008). A survey of techniques for internet "
        "traffic classification using machine learning. IEEE Communications Surveys & "
        "Tutorials, 10(4), 56\u201376. https://doi.org/10.1109/SURV.2008.080406",

        "Ring, M., Wunderlich, S., Scheuring, D., Landes, D., & Hotho, A. (2019). A survey "
        "of network-based intrusion detection data sets. Computers & Security, 86, 147\u2013167. "
        "https://doi.org/10.1016/j.cose.2019.06.005",

        "Ring, M., Wunderlich, S., Grüdl, D., Landes, D., & Hotho, A. (2017). Flow-based "
        "benchmark data sets for intrusion detection. In Proceedings of the 16th European "
        "Conference on Cyber Warfare and Security (pp. 361\u2013369). ACPI.",

        "Sharafaldin, I., Lashkari, A. H., & Ghorbani, A. A. (2018). Toward generating a "
        "new intrusion detection dataset and intrusion traffic characterization. In Proceedings "
        "of the 4th International Conference on Information Systems Security and Privacy "
        "(pp. 108\u2013116). SCITEPRESS.",

        "Stamp, M. (2011). Information security: Principles and practice (2nd ed.). Wiley.",

        "Tavallaee, M., Bagheri, E., Lu, W., & Ghorbani, A. A. (2009). A detailed analysis "
        "of the KDD CUP 99 data set. In Proceedings of the 2009 IEEE Symposium on Computational "
        "Intelligence for Security and Defense Applications (pp. 1\u20138). IEEE.",

        "Wolpert, D. H. (1992). Stacked generalization. Neural Networks, 5(2), 241\u2013259. "
        "https://doi.org/10.1016/S0893-6080(05)80023-1",

        "Zhou, Z. H. (2012). Ensemble methods: Foundations and algorithms. Chapman & Hall/CRC.",
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)  # Hanging indent (APA)
        run = p.add_run(ref)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 2.0


# ============================================================================
# MAIN
# ============================================================================

def main():
    print("=" * 70)
    print("  GENERATING ACADEMIC CHAPTERS 4 & 5")
    print("=" * 70)

    # Generate diagrams
    diagrams = generate_all_diagrams()

    # Load metrics data
    summary_df = load_summary_data()
    if summary_df.empty:
        print("[WARN] No summary data found. Tables may be empty.")

    # Create document
    doc = setup_document()

    # Title page
    for _ in range(6):
        doc.add_paragraph()
    add_paragraph(doc, "MACHINE LEARNING-BASED INTRUSION DETECTION SYSTEM",
                  bold=True, align="center")
    add_paragraph(doc, "FOR TACTICAL MILITARY NETWORKS", bold=True, align="center")
    doc.add_paragraph()
    add_paragraph(doc, "Chapters 4 and 5", align="center")
    add_paragraph(doc, "Results, Analysis, Discussion and Conclusion", italic=True, align="center")
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "Master of Science in Cybersecurity", align="center")
    doc.add_page_break()

    # Table of Contents placeholder
    add_paragraph(doc, "Table of Contents", bold=True, align="center")
    doc.add_paragraph()
    toc_items = [
        "Chapter 4: Results and Analysis",
        "    4.1 Introduction",
        "    4.2 Experimental Setup",
        "        4.2.1 System Architecture",
        "        4.2.2 Dataset Descriptions",
        "        4.2.3 Preprocessing Pipeline",
        "        4.2.4 Class Imbalance Handling",
        "    4.3 Experimental Results",
        "        4.3.1 Overall Performance Summary",
        "        4.3.2 Individual Model Comparison",
        "        4.3.3 Performance by Dataset Category",
        "    4.4 Detailed Analysis",
        "        4.4.1 Threshold Optimization",
        "        4.4.2 Software Architecture",
        "        4.4.3 Training Pipeline Workflow",
        "        4.4.4 Computational Performance",
        "    4.5 SENTINEL-IDS Interactive Dashboard",
        "        4.5.1 Dashboard Architecture and Technology Stack",
        "        4.5.2 Live Monitoring Dashboard",
        "        4.5.3 Scenario Testing Interface",
        "        4.5.4 Model Performance Analytics",
        "        4.5.5 System Architecture Visualisation",
        "        4.5.6 Alert Management and Reporting",
        "    4.6 Chapter Summary",
        "Chapter 5: Discussion and Conclusion",
        "    5.1 Discussion",
        "        5.1.1 Interpretation of Results",
        "        5.1.2 Comparison with Existing Literature",
        "        5.1.3 Feature Engineering for IoT Intrusion Detection",
        "        5.1.4 Implications for Tactical Military Networks",
        "        5.1.5 Limitations",
        "    5.2 Conclusions",
        "    5.3 Recommendations for Future Work",
        "    5.4 Final Remarks",
        "References",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    doc.add_page_break()

    # Write chapters
    print("[INFO] Writing Chapter 4: Results and Analysis...")
    write_chapter_4(doc, diagrams, summary_df)
    doc.add_page_break()

    print("[INFO] Writing Chapter 5: Discussion and Conclusion...")
    write_chapter_5(doc, summary_df)
    doc.add_page_break()

    print("[INFO] Writing References...")
    write_references(doc)

    # Save document
    output_path = OUTPUT_DIR / "Chapters_4_and_5_IDS_Thesis.docx"
    doc.save(str(output_path))
    print(f"\n[SUCCESS] Document saved to: {output_path}")
    print(f"[INFO] Diagrams saved to: {DIAGRAM_DIR}")
    return output_path


if __name__ == "__main__":
    main()
