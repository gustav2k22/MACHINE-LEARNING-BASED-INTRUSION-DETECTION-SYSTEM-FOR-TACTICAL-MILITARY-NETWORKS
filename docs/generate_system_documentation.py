#!/usr/bin/env python3
"""
Generate comprehensive system documentation for the SENTINEL-IDS project.
Creates a detailed Word document (.docx) explaining the entire system
from data preprocessing through model training to dashboard implementation.
Written for beginners with step-by-step explanations.

Usage:
    python docs/generate_system_documentation.py
"""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

OUTPUT_DIR = PROJECT_ROOT / "docs" / "output"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


def setup_styles(doc):
    """Configure document styles with Times New Roman font."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    for level in range(1, 5):
        style_name = f'Heading {level}'
        if style_name in doc.styles:
            heading_style = doc.styles[style_name]
            heading_style.font.name = 'Times New Roman'
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(0, 51, 102)
            if level == 1:
                heading_style.font.size = Pt(18)
            elif level == 2:
                heading_style.font.size = Pt(16)
            elif level == 3:
                heading_style.font.size = Pt(14)
            else:
                heading_style.font.size = Pt(12)


def add_paragraph(doc, text, bold=False, indent=False, style='Normal'):
    """Add a formatted paragraph to the document."""
    p = doc.add_paragraph(style=style)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_code_block(doc, code, language="python"):
    """Add a formatted code block to the document."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 100, 0)
    return p


def add_bullet_list(doc, items):
    """Add a bulleted list to the document."""
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)


def add_numbered_list(doc, items):
    """Add a numbered list to the document."""
    for item in items:
        p = doc.add_paragraph(item, style='List Number')
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)


def write_title_page(doc):
    """Write the title page."""
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SENTINEL-IDS")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Machine Learning-Based Intrusion Detection System\nfor Tactical Military Networks")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run("Complete System Documentation")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_paragraph()
    
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run("A Beginner-Friendly Technical Guide")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.italic = True
    
    doc.add_page_break()


def write_table_of_contents(doc):
    """Write the table of contents."""
    doc.add_heading("Table of Contents", level=1)
    
    toc_entries = [
        "1. Introduction",
        "    1.1 What is SENTINEL-IDS?",
        "    1.2 System Overview",
        "    1.3 Technology Stack",
        "2. Project Structure",
        "    2.1 Directory Layout",
        "    2.2 Key Files Explained",
        "3. Data Preprocessing",
        "    3.1 Understanding the Datasets",
        "    3.2 How Data Loading Works",
        "    3.3 Data Cleaning Process",
        "    3.4 Label Encoding",
        "    3.5 Handling Missing Values",
        "4. Feature Engineering",
        "    4.1 What is Feature Engineering?",
        "    4.2 Removing Constant Features",
        "    4.3 Correlation Analysis",
        "    4.4 Feature Scaling",
        "    4.5 Handling Class Imbalance with SMOTE",
        "5. Machine Learning Models",
        "    5.1 Understanding Ensemble Learning",
        "    5.2 Random Forest Classifier",
        "    5.3 XGBoost Classifier",
        "    5.4 LightGBM Classifier",
        "    5.5 Multi-Layer Perceptron (MLP)",
        "    5.6 Stacking Ensemble",
        "6. Training Pipeline",
        "    6.1 Pipeline Overview",
        "    6.2 Step-by-Step Training Process",
        "    6.3 Cross-Validation",
        "    6.4 Model Evaluation Metrics",
        "    6.5 Threshold Optimization",
        "7. SENTINEL-IDS Dashboard",
        "    7.1 Dashboard Architecture",
        "    7.2 React Components Explained",
        "    7.3 Live Monitoring Page",
        "    7.4 Scenario Testing Page",
        "    7.5 Model Performance Page",
        "    7.6 Architecture Visualization Page",
        "    7.7 Alerts Management Page",
        "    7.8 Theme System",
        "8. Running the System",
        "    8.1 Prerequisites",
        "    8.2 Installation Steps",
        "    8.3 Training Models",
        "    8.4 Starting the Dashboard",
        "9. Troubleshooting",
        "10. Glossary",
    ]
    
    for entry in toc_entries:
        p = doc.add_paragraph(entry)
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
    
    doc.add_page_break()


def write_introduction(doc):
    """Write the introduction chapter."""
    doc.add_heading("1. Introduction", level=1)
    
    doc.add_heading("1.1 What is SENTINEL-IDS?", level=2)
    add_paragraph(doc, 
        "SENTINEL-IDS is a sophisticated Intrusion Detection System (IDS) that uses machine learning "
        "to identify cyber attacks in network traffic. Think of it as a smart security guard for computer "
        "networks - it watches all the data flowing through a network and can tell the difference between "
        "normal, legitimate traffic and potentially malicious attacks.", indent=True)
    
    add_paragraph(doc,
        "The system was designed specifically for tactical military networks, which have unique requirements: "
        "they need to be extremely accurate (missing an attack could be catastrophic), very fast (decisions "
        "must be made in milliseconds), and able to work with limited computing resources (military equipment "
        "in the field may not have powerful computers).", indent=True)
    
    doc.add_heading("1.2 System Overview", level=2)
    add_paragraph(doc,
        "The SENTINEL-IDS system consists of three main components:", indent=True)
    
    add_numbered_list(doc, [
        "Data Preprocessing Pipeline: Takes raw network traffic data, cleans it, and transforms it into a format that machine learning models can understand.",
        "Machine Learning Engine: Uses an ensemble of five different models working together to classify network traffic as either 'Normal' or 'Attack'.",
        "Interactive Dashboard: A web-based interface that allows security analysts to monitor the network in real-time, test different scenarios, and manage alerts.",
    ])
    
    doc.add_heading("1.3 Technology Stack", level=2)
    add_paragraph(doc, "The system is built using these technologies:", indent=True)
    
    add_paragraph(doc, "Backend (Python):", bold=True)
    add_bullet_list(doc, [
        "Python 3.9+: The main programming language",
        "scikit-learn: Machine learning library for Random Forest and MLP",
        "XGBoost: Gradient boosting library",
        "LightGBM: Fast gradient boosting library",
        "pandas: Data manipulation and analysis",
        "numpy: Numerical computing",
        "imbalanced-learn: For handling imbalanced datasets (SMOTE)",
    ])
    
    add_paragraph(doc, "Frontend (JavaScript):", bold=True)
    add_bullet_list(doc, [
        "React 19: User interface library",
        "Vite 7: Build tool and development server",
        "Tailwind CSS 4: Utility-first CSS framework",
        "Recharts: Charting library for data visualization",
        "Lucide React: Icon library",
    ])
    
    doc.add_page_break()


def write_project_structure(doc):
    """Write the project structure chapter."""
    doc.add_heading("2. Project Structure", level=1)
    
    doc.add_heading("2.1 Directory Layout", level=2)
    add_paragraph(doc,
        "Understanding the project structure is essential for navigating the codebase. Here's how "
        "the project is organized:", indent=True)
    
    structure = """
SENTINEL-IDS/
├── datasets/                    # Raw training data (8 benchmark datasets)
├── models/                      # Saved trained models (.joblib files)
├── outputs/
│   ├── reports/                 # Metrics reports (JSON, CSV)
│   └── figures/                 # Generated charts and confusion matrices
├── src/                         # Python source code
│   ├── config.py                # Configuration settings
│   ├── train_pipeline.py        # Main training script
│   ├── data/
│   │   ├── preprocessor.py      # Dataset loaders
│   │   └── feature_engineering.py
│   ├── models/
│   │   └── ensemble.py          # ML model definitions
│   ├── evaluation/
│   │   ├── metrics.py           # Evaluation metrics
│   │   └── visualizations.py    # Chart generation
│   └── monitoring/
│       └── dashboard.py         # Streamlit dashboard
├── sentinel-dashboard/          # React dashboard
│   ├── src/
│   │   ├── App.jsx              # Main application
│   │   ├── components/          # Reusable UI components
│   │   ├── pages/               # Page components
│   │   ├── hooks/               # Custom React hooks
│   │   ├── context/             # React context providers
│   │   └── data/                # Constants and generators
│   └── package.json             # Node.js dependencies
├── docs/                        # Documentation
└── requirements.txt             # Python dependencies
"""
    add_code_block(doc, structure, "text")
    
    doc.add_heading("2.2 Key Files Explained", level=2)
    
    add_paragraph(doc, "src/config.py", bold=True)
    add_paragraph(doc,
        "This is the central configuration file. It defines important settings like file paths, "
        "dataset locations, and training parameters. When you need to change how the system behaves "
        "(like changing the train/test split ratio), you modify this file.", indent=True)
    
    add_paragraph(doc, "src/train_pipeline.py", bold=True)
    add_paragraph(doc,
        "This is the main entry point for training models. When you run this script, it loads all "
        "datasets, preprocesses them, trains all five models, evaluates their performance, and saves "
        "the results. Think of it as the 'conductor' that orchestrates the entire training process.", indent=True)
    
    add_paragraph(doc, "src/data/preprocessor.py", bold=True)
    add_paragraph(doc,
        "Contains specialized loading functions for each of the 8 datasets. Each dataset has a unique "
        "file format and column structure, so this file handles converting them all into a standardized "
        "format that the models can understand.", indent=True)
    
    add_paragraph(doc, "src/models/ensemble.py", bold=True)
    add_paragraph(doc,
        "Defines all five machine learning models and how they work together in the stacking ensemble. "
        "This is where the 'intelligence' of the system is defined.", indent=True)
    
    doc.add_page_break()


def write_preprocessing(doc):
    """Write the data preprocessing chapter."""
    doc.add_heading("3. Data Preprocessing", level=1)
    
    add_paragraph(doc,
        "Data preprocessing is the crucial first step in any machine learning project. Raw data from "
        "the real world is messy, inconsistent, and not in a format that algorithms can understand. "
        "Preprocessing transforms this raw data into clean, numerical values that models can learn from.", indent=True)
    
    doc.add_heading("3.1 Understanding the Datasets", level=2)
    add_paragraph(doc,
        "SENTINEL-IDS is trained on 8 different intrusion detection datasets, each representing different "
        "types of network traffic and attack scenarios:", indent=True)
    
    add_numbered_list(doc, [
        "NSL-KDD: An improved version of the famous KDD Cup 1999 dataset, containing traditional network attacks like DoS, Probe, R2L, and U2R.",
        "KDDCup99: The original benchmark dataset for intrusion detection research with ~5 million connection records.",
        "Kaggle Network Intrusion Dataset: A community-contributed dataset with modern network traffic patterns.",
        "CIC-DDoS2019: Focuses specifically on Distributed Denial of Service attacks with flow-based features.",
        "CIDDS-001: Contains flow-based network data from a simulated small business environment.",
        "DS2OS: An IoT (Internet of Things) dataset from smart home environments - very different from traditional network data.",
        "LUFlow: Flow-based network traffic from Lancaster University's network.",
        "NetworkLogs: Application-level network logs with a focus on login attempts and authentication.",
    ])
    
    doc.add_heading("3.2 How Data Loading Works", level=2)
    add_paragraph(doc,
        "Each dataset has its own loader function in preprocessor.py. Here's a simplified example of "
        "how the NSL-KDD dataset is loaded:", indent=True)
    
    code = '''def load_nsl_kdd():
    """Load and preprocess the NSL-KDD dataset."""
    # 1. Read the CSV file
    df = pd.read_csv(path / "KDDTrain+.txt", header=None)
    
    # 2. Assign column names (the file has no headers)
    df.columns = ['duration', 'protocol_type', 'service', 'flag', 
                  'src_bytes', 'dst_bytes', ...]
    
    # 3. Convert labels to binary (0=Normal, 1=Attack)
    df['label'] = df['label'].apply(
        lambda x: 0 if x == 'normal' else 1
    )
    
    # 4. Encode categorical columns as numbers
    for col in ['protocol_type', 'service', 'flag']:
        df[col] = LabelEncoder().fit_transform(df[col])
    
    return df'''
    add_code_block(doc, code)
    
    doc.add_heading("3.3 Data Cleaning Process", level=2)
    add_paragraph(doc,
        "Data cleaning involves fixing issues in the raw data:", indent=True)
    
    add_bullet_list(doc, [
        "Removing duplicate rows: Sometimes datasets contain exact duplicates that would bias the model.",
        "Handling infinity values: Some calculations (like ratios) can produce infinite values, which we replace with large numbers or zeros.",
        "Converting data types: Ensuring all numerical columns are actually stored as numbers, not text.",
        "Standardizing column names: Making column names consistent (lowercase, no spaces) across all datasets.",
    ])
    
    doc.add_heading("3.4 Label Encoding", level=2)
    add_paragraph(doc,
        "Machine learning models only understand numbers, not text. Label encoding converts text "
        "categories into numerical values. For example:", indent=True)
    
    add_code_block(doc, '''# Before encoding:
protocol_type: ['tcp', 'udp', 'icmp', 'tcp', 'udp']

# After encoding:
protocol_type: [2, 1, 0, 2, 1]  # alphabetically: icmp=0, tcp=1, udp=2''')
    
    add_paragraph(doc,
        "The LabelEncoder from scikit-learn automatically assigns a unique number to each unique "
        "text value. This is done for columns like protocol_type, service, and flag.", indent=True)
    
    doc.add_heading("3.5 Handling Missing Values", level=2)
    add_paragraph(doc,
        "Missing values (NaN - 'Not a Number') can break machine learning algorithms. Our approach:", indent=True)
    
    add_numbered_list(doc, [
        "First, we identify which columns have missing values using pandas: df.isnull().sum()",
        "For numerical columns, we fill missing values with 0 (conservative approach for network data).",
        "For categorical columns, we fill with the most common value (mode).",
        "Any rows that still have missing values after these steps are dropped.",
    ])
    
    add_code_block(doc, '''# Fill missing numerical values with 0
df = df.fillna(0)

# Or fill with column mean (alternative approach)
df = df.fillna(df.mean())''')
    
    doc.add_page_break()


def write_feature_engineering(doc):
    """Write the feature engineering chapter."""
    doc.add_heading("4. Feature Engineering", level=1)
    
    doc.add_heading("4.1 What is Feature Engineering?", level=2)
    add_paragraph(doc,
        "Feature engineering is the process of using domain knowledge to create new features or transform "
        "existing ones to improve model performance. Good feature engineering can make the difference between "
        "a mediocre model and an excellent one.", indent=True)
    
    add_paragraph(doc,
        "In SENTINEL-IDS, feature engineering happens in src/data/feature_engineering.py and includes "
        "several important steps.", indent=True)
    
    doc.add_heading("4.2 Removing Constant Features", level=2)
    add_paragraph(doc,
        "A constant feature is a column where every single row has the same value. These features "
        "provide no information for distinguishing between normal and attack traffic, so we remove them.", indent=True)
    
    add_code_block(doc, '''from sklearn.feature_selection import VarianceThreshold

# Remove features with zero variance (constant features)
selector = VarianceThreshold(threshold=0)
X_filtered = selector.fit_transform(X)

# Example: If a column 'land' has value 0 for all 100,000 rows,
# it will be removed because it doesn't help distinguish attacks.''')
    
    doc.add_heading("4.3 Correlation Analysis", level=2)
    add_paragraph(doc,
        "Highly correlated features provide redundant information. If two features always move together "
        "(correlation > 0.95), we keep only one of them. This reduces computational cost and can improve "
        "model performance.", indent=True)
    
    add_code_block(doc, '''def remove_correlated_features(X, threshold=0.95):
    """Remove features with correlation above threshold."""
    # Calculate correlation matrix
    corr_matrix = X.corr().abs()
    
    # Find pairs with high correlation
    upper = corr_matrix.where(
        np.triu(np.ones(corr_matrix.shape), k=1).astype(bool)
    )
    
    # Identify columns to drop
    to_drop = [col for col in upper.columns 
               if any(upper[col] > threshold)]
    
    return X.drop(columns=to_drop)''')
    
    doc.add_heading("4.4 Feature Scaling", level=2)
    add_paragraph(doc,
        "Different features can have vastly different ranges. For example, 'duration' might range from "
        "0 to 60000, while 'land' might only be 0 or 1. This can cause problems for many machine learning "
        "algorithms that assume features are on similar scales.", indent=True)
    
    add_paragraph(doc,
        "We use RobustScaler, which scales features based on the median and interquartile range. This is "
        "more robust to outliers than standard scaling:", indent=True)
    
    add_code_block(doc, '''from sklearn.preprocessing import RobustScaler

scaler = RobustScaler()
X_scaled = scaler.fit_transform(X)

# RobustScaler formula for each value:
# scaled_value = (value - median) / (Q3 - Q1)
# where Q1 = 25th percentile, Q3 = 75th percentile''')
    
    doc.add_heading("4.5 Handling Class Imbalance with SMOTE", level=2)
    add_paragraph(doc,
        "Class imbalance occurs when one class (e.g., 'Normal') has many more samples than another "
        "(e.g., 'Attack'). This is common in intrusion detection - most traffic is normal, with only "
        "a small percentage being attacks.", indent=True)
    
    add_paragraph(doc,
        "SMOTE (Synthetic Minority Over-sampling Technique) creates synthetic examples of the minority "
        "class to balance the dataset:", indent=True)
    
    add_code_block(doc, '''from imblearn.over_sampling import SMOTE

# Before SMOTE:
# Normal: 95,000 samples (95%)
# Attack:  5,000 samples (5%)

smote = SMOTE(random_state=42)
X_balanced, y_balanced = smote.fit_resample(X, y)

# After SMOTE:
# Normal: 95,000 samples (50%)
# Attack: 95,000 samples (50%) <- synthetic samples created''')
    
    add_paragraph(doc,
        "How SMOTE works: For each minority class sample, SMOTE finds its k nearest neighbors (default k=5) "
        "and creates new synthetic samples along the line segments connecting the original sample to its neighbors.", indent=True)
    
    add_paragraph(doc,
        "In SENTINEL-IDS, we use a tiered SMOTE strategy based on the imbalance ratio:", indent=True)
    
    add_bullet_list(doc, [
        "Extreme imbalance (>10:1 ratio): Conservative oversampling to 3x the minority class size",
        "Moderate imbalance (3:1 to 10:1): SMOTE to reach 50% of majority class size",
        "Near-balanced (<3:1): Skip SMOTE entirely to avoid overfitting",
    ])
    
    doc.add_page_break()


def write_ml_models(doc):
    """Write the machine learning models chapter."""
    doc.add_heading("5. Machine Learning Models", level=1)
    
    doc.add_heading("5.1 Understanding Ensemble Learning", level=2)
    add_paragraph(doc,
        "Instead of relying on a single model, SENTINEL-IDS uses an ensemble of five models working "
        "together. This is like asking five different experts for their opinions and then combining "
        "them - the combined wisdom is usually better than any single expert.", indent=True)
    
    add_paragraph(doc,
        "The system uses 'stacking', where the outputs of four base models are fed into a fifth "
        "'meta-learner' model that makes the final decision.", indent=True)
    
    doc.add_heading("5.2 Random Forest Classifier", level=2)
    add_paragraph(doc,
        "Random Forest is a collection of decision trees, where each tree 'votes' on the classification. "
        "The final prediction is the majority vote of all trees.", indent=True)
    
    add_paragraph(doc, "How it works:", bold=True)
    add_numbered_list(doc, [
        "Training: Create 300 decision trees, each trained on a random subset of the data and features.",
        "Prediction: Each tree makes a prediction. If 200 trees say 'Attack' and 100 say 'Normal', the final prediction is 'Attack'.",
        "Probability: The confidence is the percentage of trees agreeing (e.g., 200/300 = 66.7% confidence).",
    ])
    
    add_code_block(doc, '''RandomForestClassifier(
    n_estimators=300,        # Number of trees
    max_depth=25,            # Maximum tree depth
    class_weight='balanced', # Handle imbalanced classes
    random_state=42,         # For reproducibility
    n_jobs=-1                # Use all CPU cores
)''')
    
    doc.add_heading("5.3 XGBoost Classifier", level=2)
    add_paragraph(doc,
        "XGBoost (eXtreme Gradient Boosting) builds trees sequentially, where each new tree tries to "
        "correct the mistakes of the previous trees. It's like a team where each member focuses on "
        "fixing what the previous members got wrong.", indent=True)
    
    add_code_block(doc, '''XGBClassifier(
    n_estimators=300,          # Number of boosting rounds
    max_depth=10,              # Tree depth
    learning_rate=0.1,         # How much each tree contributes
    scale_pos_weight=ratio,    # Handle class imbalance
    use_label_encoder=False,
    eval_metric='logloss'
)''')
    
    add_paragraph(doc,
        "The scale_pos_weight parameter is dynamically calculated based on the class imbalance ratio "
        "of each dataset: scale_pos_weight = (number of negatives) / (number of positives)", indent=True)
    
    doc.add_heading("5.4 LightGBM Classifier", level=2)
    add_paragraph(doc,
        "LightGBM is similar to XGBoost but uses a different tree-building algorithm (leaf-wise instead "
        "of level-wise) that makes it faster while often achieving similar accuracy.", indent=True)
    
    add_code_block(doc, '''LGBMClassifier(
    n_estimators=300,
    max_depth=15,
    learning_rate=0.1,
    is_unbalance=True,  # Built-in imbalance handling
    verbose=-1          # Suppress warnings
)''')
    
    doc.add_heading("5.5 Multi-Layer Perceptron (MLP)", level=2)
    add_paragraph(doc,
        "MLP is a neural network with multiple layers of neurons. Unlike tree-based models, it can learn "
        "complex non-linear relationships in the data.", indent=True)
    
    add_paragraph(doc, "Architecture:", bold=True)
    add_bullet_list(doc, [
        "Input layer: Receives the feature vector (e.g., 41 features for NSL-KDD)",
        "Hidden layer 1: 128 neurons with ReLU activation",
        "Hidden layer 2: 64 neurons with ReLU activation", 
        "Hidden layer 3: 32 neurons with ReLU activation",
        "Output layer: 2 neurons (Normal or Attack) with softmax activation",
    ])
    
    add_code_block(doc, '''MLPClassifier(
    hidden_layer_sizes=(128, 64, 32),
    activation='relu',
    solver='adam',
    learning_rate='adaptive',
    early_stopping=True,    # Stop when validation loss stops improving
    max_iter=200
)''')
    
    doc.add_heading("5.6 Stacking Ensemble", level=2)
    add_paragraph(doc,
        "The stacking ensemble combines all four models using a meta-learner (Logistic Regression) "
        "that learns the optimal way to combine their predictions.", indent=True)
    
    add_paragraph(doc, "How stacking works:", bold=True)
    add_numbered_list(doc, [
        "Train each base model on the training data using 5-fold cross-validation.",
        "For each fold, collect the predictions (probabilities) from each base model.",
        "These predictions become the input features for the meta-learner.",
        "The meta-learner learns which models to trust in different situations.",
        "Final prediction: Base models make predictions → Meta-learner combines them → Output.",
    ])
    
    add_code_block(doc, '''from sklearn.ensemble import StackingClassifier

stacking_ensemble = StackingClassifier(
    estimators=[
        ('rf', random_forest),
        ('xgb', xgboost),
        ('lgb', lightgbm),
        ('mlp', mlp)
    ],
    final_estimator=LogisticRegression(max_iter=1000),
    cv=5  # 5-fold cross-validation for generating meta-features
)''')
    
    doc.add_page_break()


def write_training_pipeline(doc):
    """Write the training pipeline chapter."""
    doc.add_heading("6. Training Pipeline", level=1)
    
    doc.add_heading("6.1 Pipeline Overview", level=2)
    add_paragraph(doc,
        "The training pipeline in train_pipeline.py orchestrates the entire process from raw data to "
        "saved models. Here's the high-level flow:", indent=True)
    
    add_code_block(doc, '''For each dataset:
    1. Load raw data → preprocessor.py
    2. Clean and encode → preprocessor.py
    3. Feature engineering → feature_engineering.py
    4. Split into train/test (80/20)
    5. Train all 5 models → ensemble.py
    6. Evaluate on test set → metrics.py
    7. Optimize thresholds if needed → metrics.py
    8. Save models and results''')
    
    doc.add_heading("6.2 Step-by-Step Training Process", level=2)
    
    add_paragraph(doc, "Step 1: Load Dataset", bold=True)
    add_code_block(doc, '''df = load_dataset("NSL_KDD")
# Returns a pandas DataFrame with features and 'label' column''')
    
    add_paragraph(doc, "Step 2: Prepare Features", bold=True)
    add_code_block(doc, '''X_train, X_test, y_train, y_test, scaler, feature_names = prepare_dataset(df)
# Returns scaled, balanced training data and unmodified test data''')
    
    add_paragraph(doc, "Step 3: Build Models", bold=True)
    add_code_block(doc, '''imbalance_ratio = len(y_train[y_train==0]) / len(y_train[y_train==1])
models = build_base_models(imbalance_ratio)
# Returns dict: {'RandomForest': rf, 'XGBoost': xgb, 'LightGBM': lgb, 'MLP': mlp}''')
    
    add_paragraph(doc, "Step 4: Train Each Model", bold=True)
    add_code_block(doc, '''for name, model in models.items():
    model.fit(X_train, y_train)
    y_pred = model.predict(X_test)
    y_proba = model.predict_proba(X_test)[:, 1]  # Probability of attack
    metrics = compute_all_metrics(y_test, y_pred, y_proba)''')
    
    add_paragraph(doc, "Step 5: Train Stacking Ensemble", bold=True)
    add_code_block(doc, '''stacking_model = train_stacking_ensemble(models, X_train, y_train)
# Uses 5-fold CV to generate meta-features, then trains the meta-learner''')
    
    add_paragraph(doc, "Step 6: Save Results", bold=True)
    add_code_block(doc, '''joblib.dump(model, f"models/{dataset}_{model_name}.joblib")
# Also saves scaler, feature names, and optimal threshold''')
    
    doc.add_heading("6.3 Cross-Validation", level=2)
    add_paragraph(doc,
        "Cross-validation helps us estimate how well the model will perform on unseen data. We use "
        "5-fold cross-validation:", indent=True)
    
    add_numbered_list(doc, [
        "Split training data into 5 equal parts (folds).",
        "Train on 4 folds, validate on 1 fold. Repeat 5 times, each time using a different fold for validation.",
        "Average the performance across all 5 folds.",
        "This gives a more reliable estimate than a single train/test split.",
    ])
    
    doc.add_heading("6.4 Model Evaluation Metrics", level=2)
    add_paragraph(doc,
        "We evaluate models using multiple metrics because no single metric tells the whole story:", indent=True)
    
    add_bullet_list(doc, [
        "Accuracy: (Correct predictions) / (Total predictions) - Simple but can be misleading with imbalanced data.",
        "Precision: (True Attacks) / (Predicted Attacks) - Of everything we flagged as an attack, how many were real attacks?",
        "Recall: (True Attacks) / (Actual Attacks) - Of all actual attacks, how many did we catch?",
        "F1 Score: Harmonic mean of Precision and Recall - Balances both concerns.",
        "ROC-AUC: Area Under the ROC Curve - Measures ability to distinguish between classes at all thresholds.",
        "Specificity: (True Normal) / (Actual Normal) - Of all normal traffic, how much did we correctly identify?",
        "MCC: Matthews Correlation Coefficient - A balanced measure even with imbalanced classes.",
    ])
    
    doc.add_heading("6.5 Threshold Optimization", level=2)
    add_paragraph(doc,
        "By default, if a model predicts >50% probability of attack, we classify it as an attack. But "
        "this threshold can be adjusted:", indent=True)
    
    add_bullet_list(doc, [
        "Lower threshold (e.g., 30%): More sensitive, catches more attacks but also more false alarms.",
        "Higher threshold (e.g., 70%): More conservative, fewer false alarms but might miss some attacks.",
    ])
    
    add_paragraph(doc,
        "SENTINEL-IDS automatically searches for the optimal threshold that maximizes F1 score while "
        "keeping both precision and recall above 0.9:", indent=True)
    
    add_code_block(doc, '''def optimize_threshold(y_true, y_proba, target=0.9):
    best_threshold = 0.5
    best_f1 = 0
    
    for threshold in np.arange(0.01, 0.99, 0.01):
        y_pred = (y_proba >= threshold).astype(int)
        precision = precision_score(y_true, y_pred)
        recall = recall_score(y_true, y_pred)
        f1 = f1_score(y_true, y_pred)
        
        if precision >= target and recall >= target and f1 > best_f1:
            best_f1 = f1
            best_threshold = threshold
    
    return best_threshold''')
    
    doc.add_page_break()


def write_dashboard(doc):
    """Write the dashboard chapter."""
    doc.add_heading("7. SENTINEL-IDS Dashboard", level=1)
    
    add_paragraph(doc,
        "The SENTINEL-IDS Dashboard is a React-based web application that provides a user-friendly "
        "interface for interacting with the intrusion detection system. It's designed with a military "
        "aesthetic and supports both dark and light modes.", indent=True)
    
    doc.add_heading("7.1 Dashboard Architecture", level=2)
    add_paragraph(doc,
        "The dashboard follows a clean, modular architecture separating different concerns:", indent=True)
    
    add_code_block(doc, '''sentinel-dashboard/src/
├── App.jsx              # Root component with routing
├── context/
│   └── ThemeContext.jsx # Dark/light mode management
├── hooks/
│   └── useSimulation.js # Live data simulation
├── components/          # Reusable UI pieces
│   ├── Panel.jsx        # Container component
│   ├── KPICard.jsx      # Metric display card
│   ├── SeverityBadge.jsx# Alert severity indicator
│   ├── ThreatMap.jsx    # Network visualization
│   ├── Sidebar.jsx      # Navigation menu
│   └── TopBar.jsx       # Status bar
├── pages/               # Full page components
│   ├── DashboardPage.jsx
│   ├── ScenarioTestPage.jsx
│   ├── ModelPerformancePage.jsx
│   ├── ArchitecturePage.jsx
│   └── AlertsPage.jsx
└── data/                # Constants and mock data
    ├── constants.js
    ├── metrics.js
    └── generators.js''')
    
    doc.add_heading("7.2 React Components Explained", level=2)
    add_paragraph(doc,
        "React components are reusable pieces of UI. Here's a simple example of how the KPICard "
        "component works:", indent=True)
    
    add_code_block(doc, '''// KPICard.jsx - Displays a key metric with trend
export function KPICard({ icon: Icon, label, value, trend, accent }) {
    return (
        <div className="border p-4 bg-charcoal">
            <div className="flex items-center gap-2">
                <Icon size={20} color={accent} />
                <span className="text-muted text-xs">{label}</span>
            </div>
            <p className="text-2xl font-bold">{value}</p>
            <span className={trend > 0 ? "text-red" : "text-green"}>
                {trend > 0 ? "↑" : "↓"} {Math.abs(trend)}% vs last hr
            </span>
        </div>
    );
}

// Usage:
<KPICard 
    icon={Activity} 
    label="PACKETS ANALYZED (TODAY)" 
    value="128,465" 
    trend={12}
    accent="#00D4FF"
/>''')
    
    doc.add_heading("7.3 Live Monitoring Page", level=2)
    add_paragraph(doc,
        "The main dashboard page shows real-time network activity:", indent=True)
    
    add_bullet_list(doc, [
        "KPI Cards: Four cards showing packets analyzed, attacks detected, false positive rate, and system uptime.",
        "Tactical Threat Map: An animated SVG showing network nodes and connections, with attack flows highlighted in red.",
        "Threat Timeline: A 60-minute rolling chart showing the volume of normal vs. attack traffic.",
        "Attack Distribution: A pie chart breaking down attacks by type (DoS, Probe, R2L, U2R, DDoS).",
        "Live Traffic Feed: A scrolling table of recent packet classifications with IP addresses, protocols, and verdicts.",
        "Active Alerts: A sidebar showing the most recent high-priority alerts.",
    ])
    
    doc.add_heading("7.4 Scenario Testing Page", level=2)
    add_paragraph(doc,
        "This page allows analysts to test the model's detection capabilities:", indent=True)
    
    add_numbered_list(doc, [
        "Select a pre-built scenario (Normal Browsing, DoS Attack, Port Scan, Brute Force SSH) or create a custom one.",
        "The system displays the feature values for that scenario.",
        "Click 'Analyze Traffic' to get the model's prediction.",
        "View the verdict (ATTACK or NORMAL), confidence score, and per-model breakdown.",
        "See which features most influenced the decision (feature importance chart).",
    ])
    
    doc.add_heading("7.5 Model Performance Page", level=2)
    add_paragraph(doc,
        "Displays comprehensive analytics for all five models:", indent=True)
    
    add_bullet_list(doc, [
        "Model Overview Cards: Summary of each model's architecture and key metrics.",
        "Metrics Comparison Table: Side-by-side comparison of accuracy, precision, recall, F1, ROC-AUC, and MCC.",
        "Interactive Confusion Matrix: Shows true/false positives/negatives with model toggle.",
        "Threshold Optimization Chart: Visualizes how precision, recall, and F1 change at different thresholds.",
        "Dataset Comparison: Bar chart comparing F1 scores across all 8 datasets.",
    ])
    
    doc.add_heading("7.6 Architecture Visualization Page", level=2)
    add_paragraph(doc,
        "Shows visual diagrams of the system design:", indent=True)
    
    add_bullet_list(doc, [
        "4-Layer System Architecture: Data Layer → Preprocessing Layer → Model Layer → Evaluation Layer.",
        "Stacking Ensemble Flow: How base models feed into the meta-learner.",
        "Preprocessing Pipeline: The 11 steps from raw data to model-ready features.",
        "UML Class Diagram: Shows the relationships between Python classes.",
    ])
    
    doc.add_heading("7.7 Alerts Management Page", level=2)
    add_paragraph(doc,
        "Provides tools for managing and investigating alerts:", indent=True)
    
    add_bullet_list(doc, [
        "Search: Find alerts by IP address or other fields.",
        "Filter: Filter by severity (Critical, High, Medium) or attack type.",
        "Pagination: Navigate through large numbers of alerts.",
        "Export: Download alerts as CSV for external analysis.",
        "Detail View: Click an alert to see the full feature vector and recommended actions.",
    ])
    
    doc.add_heading("7.8 Theme System", level=2)
    add_paragraph(doc,
        "The dashboard supports automatic theme detection and manual switching:", indent=True)
    
    add_code_block(doc, '''// ThemeContext.jsx - Manages dark/light mode
export function ThemeProvider({ children }) {
    // Detect OS preference
    const prefersDark = window.matchMedia(
        '(prefers-color-scheme: dark)'
    ).matches;
    
    const [isDark, setIsDark] = useState(prefersDark);
    
    // Color palette changes based on theme
    const colors = isDark ? {
        navy: '#0A0E1A',
        charcoal: '#1A1F2E',
        white: '#FFFFFF',
        // ... more colors
    } : {
        navy: '#F8FAFC',
        charcoal: '#FFFFFF',
        white: '#0F172A',
        // ... light mode colors
    };
    
    return (
        <ThemeContext.Provider value={{ isDark, toggle, colors }}>
            {children}
        </ThemeContext.Provider>
    );
}''')
    
    doc.add_page_break()


def write_running_system(doc):
    """Write the running the system chapter."""
    doc.add_heading("8. Running the System", level=1)
    
    doc.add_heading("8.1 Prerequisites", level=2)
    add_paragraph(doc, "Before you begin, ensure you have:", indent=True)
    
    add_bullet_list(doc, [
        "Python 3.9 or higher: Check with python3 --version",
        "Node.js 18 or higher: Check with node --version",
        "npm 9 or higher: Check with npm --version",
        "At least 8GB of RAM (16GB recommended for training)",
        "At least 10GB of free disk space for datasets and models",
    ])
    
    doc.add_heading("8.2 Installation Steps", level=2)
    
    add_paragraph(doc, "Step 1: Clone or download the project", bold=True)
    add_code_block(doc, '''cd ~/Downloads
# If downloaded as ZIP, unzip it
# If using git: git clone <repository-url>''')
    
    add_paragraph(doc, "Step 2: Install Python dependencies", bold=True)
    add_code_block(doc, '''cd MACHINE-LEARNING-BASED-INTRUSION-DETECTION-SYSTEM-FOR-TACTICAL-MILITARY-NETWORKS
pip install -r requirements.txt

# On macOS, if XGBoost fails:
brew install libomp
pip install xgboost''')
    
    add_paragraph(doc, "Step 3: Install dashboard dependencies", bold=True)
    add_code_block(doc, '''cd sentinel-dashboard
npm install''')
    
    doc.add_heading("8.3 Training Models", level=2)
    add_paragraph(doc, "To train all models on all 8 datasets:", indent=True)
    add_code_block(doc, '''# From the project root directory
python -m src.train_pipeline

# To train on a specific dataset only:
python -m src.train_pipeline NSL_KDD''')
    
    add_paragraph(doc,
        "Training takes approximately 15-30 minutes depending on your hardware. Progress will be "
        "displayed in the terminal. When complete, models are saved to the models/ directory.", indent=True)
    
    doc.add_heading("8.4 Starting the Dashboard", level=2)
    add_paragraph(doc, "Development mode (with hot reload):", bold=True)
    add_code_block(doc, '''cd sentinel-dashboard
npm run dev
# Open http://localhost:5173 in your browser''')
    
    add_paragraph(doc, "Production build:", bold=True)
    add_code_block(doc, '''cd sentinel-dashboard
npm run build
# Compiled files are in dist/ directory

# To preview the production build:
npm run preview''')
    
    doc.add_page_break()


def write_troubleshooting(doc):
    """Write the troubleshooting chapter."""
    doc.add_heading("9. Troubleshooting", level=1)
    
    add_paragraph(doc, "XGBoost installation fails on macOS", bold=True)
    add_code_block(doc, '''# Install OpenMP library first
brew install libomp

# Then reinstall XGBoost
pip uninstall xgboost
pip install xgboost''')
    
    add_paragraph(doc, "Out of memory during training", bold=True)
    add_paragraph(doc,
        "If training crashes due to memory, reduce the sample size in src/config.py:", indent=True)
    add_code_block(doc, '''# Change this line:
MAX_SAMPLE_SIZE = 200_000  # Original

# To a smaller value:
MAX_SAMPLE_SIZE = 100_000  # Reduced''')
    
    add_paragraph(doc, "Dashboard shows 'Module not found' error", bold=True)
    add_code_block(doc, '''cd sentinel-dashboard
rm -rf node_modules
npm install''')
    
    add_paragraph(doc, "Charts not rendering in dashboard", bold=True)
    add_paragraph(doc,
        "Ensure your browser window is large enough. Recharts requires a minimum container size. "
        "Try maximizing the browser window or reducing zoom level.", indent=True)
    
    add_paragraph(doc, "Model file not found errors", bold=True)
    add_paragraph(doc,
        "If the dashboard can't find model files, ensure you've run the training pipeline first. "
        "The models/ directory should contain .joblib files after training.", indent=True)
    
    doc.add_page_break()


def write_glossary(doc):
    """Write the glossary chapter."""
    doc.add_heading("10. Glossary", level=1)
    
    terms = [
        ("Binary Classification", "A machine learning task where we categorize items into exactly two classes (e.g., Normal vs. Attack)."),
        ("Class Imbalance", "When one class has significantly more samples than another in the training data."),
        ("Cross-Validation", "A technique for evaluating models by training and testing on different subsets of the data."),
        ("Decision Tree", "A flowchart-like model that makes decisions based on asking a series of questions about the features."),
        ("Ensemble", "A method that combines multiple models to make better predictions than any single model."),
        ("Epoch", "One complete pass through the entire training dataset during neural network training."),
        ("F1 Score", "The harmonic mean of precision and recall, providing a single metric that balances both."),
        ("False Positive", "When the model incorrectly predicts an attack when the traffic was actually normal."),
        ("False Negative", "When the model fails to detect an actual attack (predicts normal when it was an attack)."),
        ("Feature", "An individual measurable property used as input to a machine learning model (e.g., packet duration, byte count)."),
        ("Feature Engineering", "The process of creating new features or transforming existing ones to improve model performance."),
        ("Gradient Boosting", "A technique that builds models sequentially, with each new model correcting the errors of the previous ones."),
        ("Hyperparameter", "A configuration setting for a model that is set before training (e.g., number of trees, learning rate)."),
        ("IDS (Intrusion Detection System)", "A system that monitors network traffic to detect malicious activity."),
        ("Label", "The target variable we're trying to predict (in our case, 0 for Normal, 1 for Attack)."),
        ("LabelEncoder", "A tool that converts categorical text values into numerical values."),
        ("Meta-Learner", "In stacking, the model that combines the predictions of the base models."),
        ("MLP (Multi-Layer Perceptron)", "A type of neural network with multiple layers of neurons."),
        ("Overfitting", "When a model learns the training data too well, including noise, and performs poorly on new data."),
        ("Precision", "Of all items predicted as positive, what fraction were actually positive."),
        ("Probability", "A number between 0 and 1 indicating how confident the model is in its prediction."),
        ("Random Forest", "An ensemble of decision trees that vote on the final prediction."),
        ("Recall", "Of all actual positives, what fraction did the model correctly identify."),
        ("ROC-AUC", "Area Under the Receiver Operating Characteristic curve; measures model's ability to distinguish classes."),
        ("SMOTE", "Synthetic Minority Over-sampling Technique; creates synthetic examples to balance classes."),
        ("Stacking", "An ensemble technique where predictions from multiple models are used as input to a meta-learner."),
        ("Threshold", "The probability cutoff for classification (e.g., if probability > 0.5, predict Attack)."),
        ("Training/Test Split", "Dividing data into a set for training the model and a set for evaluating it."),
        ("True Positive", "When the model correctly identifies an attack."),
        ("True Negative", "When the model correctly identifies normal traffic."),
        ("Variance", "How much a feature's values differ from each other; features with zero variance are constant."),
        ("XGBoost", "eXtreme Gradient Boosting; a fast and accurate gradient boosting implementation."),
    ]
    
    for term, definition in terms:
        p = doc.add_paragraph()
        run = p.add_run(f"{term}: ")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        run = p.add_run(definition)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)


def main():
    """Generate the complete system documentation."""
    print("=" * 60)
    print("  GENERATING SENTINEL-IDS SYSTEM DOCUMENTATION")
    print("=" * 60)
    
    doc = Document()
    setup_styles(doc)
    
    print("[INFO] Writing title page...")
    write_title_page(doc)
    
    print("[INFO] Writing table of contents...")
    write_table_of_contents(doc)
    
    print("[INFO] Writing Chapter 1: Introduction...")
    write_introduction(doc)
    
    print("[INFO] Writing Chapter 2: Project Structure...")
    write_project_structure(doc)
    
    print("[INFO] Writing Chapter 3: Data Preprocessing...")
    write_preprocessing(doc)
    
    print("[INFO] Writing Chapter 4: Feature Engineering...")
    write_feature_engineering(doc)
    
    print("[INFO] Writing Chapter 5: Machine Learning Models...")
    write_ml_models(doc)
    
    print("[INFO] Writing Chapter 6: Training Pipeline...")
    write_training_pipeline(doc)
    
    print("[INFO] Writing Chapter 7: Dashboard...")
    write_dashboard(doc)
    
    print("[INFO] Writing Chapter 8: Running the System...")
    write_running_system(doc)
    
    print("[INFO] Writing Chapter 9: Troubleshooting...")
    write_troubleshooting(doc)
    
    print("[INFO] Writing Chapter 10: Glossary...")
    write_glossary(doc)
    
    output_path = OUTPUT_DIR / "SENTINEL_IDS_System_Documentation.docx"
    doc.save(output_path)
    
    print()
    print(f"[SUCCESS] Documentation saved to: {output_path}")
    print("=" * 60)


if __name__ == "__main__":
    main()
